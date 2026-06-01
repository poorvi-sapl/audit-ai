"""
raw_to_training_pair/generation/gold_loader.py
================================================
Read a FILLED NPO-CX-1.1 .docx workpaper and produce a
GeneratedWorkpaper (the assistant-target half of a generation pair).

This is the reverse of workpaper_generator/renderer.py. The renderer
writes resolved field values INTO specific cell positions; this
loader reads them BACK OUT using the same manifest-driven mapping.

Phase 2.2 scope: NPO-CX-1.1 ONLY. Sibling engagement-acceptance
variants (GOV/FP/TRB) reuse 90% of this code and are added in a
follow-up once NPO is validated against real HCLLP gold docs.

Citations are emitted EMPTY by the loader — the .docx itself doesn't
carry source provenance. Use citation_linker.auto_link_citations() to
populate citations from matching ExtractedFact provenance (Decision L2).

Public API
----------
    load_filled_workpaper(filepath, workpaper_type, engagement_id="")
        → GeneratedWorkpaper
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

import yaml
from docx import Document

from raw_to_training_pair.generation.target_schema import (
    GeneratedFieldValue,
    GeneratedWorkpaper,
)

logger = logging.getLogger(__name__)

# Match renderer.py constants. If renderer.py changes these, this loader
# must update in lockstep. A future cleanup pass should hoist these into
# a shared workpaper_generator/manifest_layout.py module.
_T_HEADER = 0
_T_PART_I = 2
_T_SIGNOFF = 3
_T_PART_II = 4

_PART_I_YES_COL = 2
_PART_I_NO_COL = 3
_PART_I_COMMENT_COL = 4

_PART_II_YES_COL = 2
_PART_II_NO_COL = 3
_PART_II_NA_COL = 4
_PART_II_COMMENT_COL = 5

_MARK = "X"

# SOP "Practical Consideration" rows are notes, not questions — skip them
_SKIP_ROW_MARKERS = ("practical consideration",)

_PROJECT_ROOT = Path(__file__).parent.parent.parent
_DEFAULT_MANIFEST_PATH = _PROJECT_ROOT / "config" / "npo_cx_1_1_manifest.yaml"


# ---------------------------------------------------------------------
# Text normalization (mirrors workpaper_generator/renderer.py)
# ---------------------------------------------------------------------

def _normalize(text: str) -> str:
    text = text.lower()
    text = text.replace("-", "")
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _first_n_words(text: str, n: int = 8) -> str:
    return " ".join(_normalize(text).split()[:n])


def _is_skippable_row(row_text_norm: str) -> bool:
    return any(m in row_text_norm for m in _SKIP_ROW_MARKERS) or not row_text_norm


def _find_row(table, target_prefix: str, used_rows: set[int]) -> int | None:
    if not target_prefix:
        return None
    for ri, row in enumerate(table.rows):
        if ri in used_rows:
            continue
        c0 = row.cells[0].text
        c0_norm = _normalize(c0)
        if _is_skippable_row(c0_norm):
            continue
        c0_prefix = _first_n_words(c0)
        if c0_prefix.startswith(target_prefix):
            return ri
    return None


# ---------------------------------------------------------------------
# Cell value parsing
# ---------------------------------------------------------------------

def _has_mark(cell_text: str) -> bool:
    """True iff a Yes/No/N/A cell shows a mark via cell text.

    Detects three formats:
      1. Literal "X" (synthetic test fixture convention)
      2. Unicode Private Use Area characters U+E000-U+F8FF — this is
         what Wingdings symbols look like in modern Word when rendered
         to text (e.g., ECEstep's form uses \\uf061 for marked cells)
      3. Other common checkmark codepoints (✓ ✔ etc.)
    Real HCLLP workpapers using <w:sym> elements are handled by
    _cell_is_marked, which calls this AND inspects the cell XML.
    """
    text = cell_text.strip()
    if not text:
        return False
    if text.upper() == _MARK:
        return True
    # Check for any Unicode Private Use Area character (Wingdings ends
    # up here when rendered to text in modern Word documents).
    for ch in text:
        cp = ord(ch)
        if 0xE000 <= cp <= 0xF8FF:
            return True
        # Common explicit checkmark codepoints
        if ch in ("✓", "✔", "☑", "■", "●"):
            return True
    return False


# Wingdings character codes used by the real HCLLP NPO-CX-1.1 forms
# for Yes/No/N/A cell marks (rendered as <w:sym w:font="Wingdings"
# w:char="..."/> elements). The check character F0FC is the most
# common "checked" mark; F0FB and F0FE are nearby checkbox glyphs.
_WINGDINGS_CHECK_CHARS: frozenset[str] = frozenset({
    "F0FC", "F0FB", "F0FE", "F0FD",
    "f0fc", "f0fb", "f0fe", "f0fd",
})

# XML namespace for WordprocessingML
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _cell_is_marked(cell) -> bool:
    """True iff the cell contains a Yes/No/N/A mark.

    Real HCLLP workpapers use Wingdings symbol marks rendered via
    <w:sym w:font="Wingdings" w:char="F0FC"/> (or similar). The
    python-docx `.text` attribute does NOT surface symbol runs, so a
    text-based check misses these. Falls back to text "X" detection
    for backward compatibility with synthetic test fixtures.
    """
    # Fast path: legacy "X" text check (synthetic test fixture format)
    if _has_mark(cell.text):
        return True
    # Real-form path: look for <w:sym> elements in the cell's XML
    try:
        sym_elements = cell._tc.findall(f".//{{{_W_NS}}}sym")
    except Exception:
        return False
    if not sym_elements:
        return False
    for sym in sym_elements:
        char = sym.get(f"{{{_W_NS}}}char", "")
        if char in _WINGDINGS_CHECK_CHARS:
            return True
    # Also: if there's any sym at all in the marks column, treat as marked.
    # The Wingdings character set has many checkmark / box variants and
    # different HCLLP templates may use different codes. A sym element
    # in a Yes/No/N/A column basically always means "the auditor put a
    # mark here" — we treat it as marked.
    return True


def _strip_label_prefix(cell_text: str, label: str) -> str | None:
    """Header cells have the form 'Label: value'. Return the value
    portion, or None if the cell is empty after stripping.

    Real workpaper header cells sometimes contain multiple paragraphs
    (e.g., 'Heffernan Foundation\\nEngagement Date: 12/04/2025'). We
    take only the FIRST non-empty line as the atomic field value;
    additional paragraphs are unrelated content that doesn't belong
    in this field.
    """
    text = cell_text.strip()
    prefix = f"{label}:"
    if text.startswith(prefix):
        text = text[len(prefix):].strip()
    # Take only the first non-empty line — drop any trailing
    # secondary content like "Engagement Date: ..."
    first_line = next(
        (ln.strip() for ln in text.splitlines() if ln.strip()), "",
    )
    return first_line or None


# ---------------------------------------------------------------------
# Per-table loaders
# ---------------------------------------------------------------------

def _load_header(table) -> dict[str, GeneratedFieldValue]:
    """Header table cells contain 'Label: value' strings written by
    renderer._render_header. Reverse that to extract just the values."""
    fields: dict[str, GeneratedFieldValue] = {}

    label_map = [
        (0, 0, "Organization",                        "organization_name"),
        (0, 1, "Statement of Financial Position Date", "financial_position_date"),
        (1, 0, "Completed by",                         "completed_by"),
        (1, 1, "Date",                                 "completion_date"),
    ]
    for r, c, label, field_id in label_map:
        try:
            cell_text = table.rows[r].cells[c].text
        except IndexError:
            fields[field_id] = GeneratedFieldValue(value=None)
            continue
        fields[field_id] = GeneratedFieldValue(
            value=_strip_label_prefix(cell_text, label),
        )
    return fields


def _three_state_boolean_from_row(
    cells: list, yes_col: int, no_col: int,
    na_col: int | None = None,
) -> bool | None:
    """Read a Yes/No (or Yes/No/N/A) row. Returns True for Yes,
    False for No, None for N/A or unmarked."""
    if _cell_is_marked(cells[yes_col]):
        return True
    if _cell_is_marked(cells[no_col]):
        return False
    # na_col present → None (the registry models N/A as None for booleans)
    return None


def _comment_value(cells: list, comment_col: int) -> str | None:
    text = cells[comment_col].text.strip() if comment_col < len(cells) else ""
    # Strip the standard label prefix if it's left over from the template
    text = text.strip()
    return text or None


def _process_part_i_question(
    table, q_def: dict, fields_out: dict[str, GeneratedFieldValue],
    used_rows: set[int],
) -> None:
    """Process one Part I question definition (single field or sub-field).

    Dispatches by input_type to write 1 or 2 entries into fields_out
    (hybrid yes_no_with_X fields produce two entries: boolean + suffix).
    """
    field_id = q_def["field_id"]
    question = q_def.get("question", "")
    input_type = q_def.get("input_type", "")

    if not question:
        # Defensive: parent fields without a 'question' (rare) get a null entry
        fields_out[field_id] = GeneratedFieldValue(value=None)
        return

    prefix = _first_n_words(question)
    row_idx = _find_row(table, prefix, used_rows)
    if row_idx is None:
        # Row not found in template — emit null. Loader logs nothing here
        # because it's expected for some manifest variants (e.g., recurring-
        # vs-initial engagement skips Part II).
        fields_out[field_id] = GeneratedFieldValue(value=None)
        return
    used_rows.add(row_idx)

    cells = table.rows[row_idx].cells

    if input_type == "yes_no":
        fields_out[field_id] = GeneratedFieldValue(
            value=_three_state_boolean_from_row(
                cells, _PART_I_YES_COL, _PART_I_NO_COL,
            ),
        )
    elif input_type == "dropdown":
        fields_out[field_id] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    elif input_type == "text_prefill":
        fields_out[field_id] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    elif input_type == "yes_no_with_specification":
        bool_val = _three_state_boolean_from_row(
            cells, _PART_I_YES_COL, _PART_I_NO_COL,
        )
        fields_out[field_id] = GeneratedFieldValue(value=bool_val)
        fields_out[f"{field_id}_specification"] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    elif input_type == "yes_no_with_text":
        bool_val = _three_state_boolean_from_row(
            cells, _PART_I_YES_COL, _PART_I_NO_COL,
        )
        fields_out[field_id] = GeneratedFieldValue(value=bool_val)
        fields_out[f"{field_id}_text"] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    elif input_type == "yes_no_with_reference":
        bool_val = _three_state_boolean_from_row(
            cells, _PART_I_YES_COL, _PART_I_NO_COL,
        )
        fields_out[field_id] = GeneratedFieldValue(value=bool_val)
        fields_out[f"{field_id}_reference"] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    elif input_type == "yes_no_with_remark_on_yes":
        # q12 — boolean answer; if Yes, the comment cell holds an
        # explanatory remark. We store both as separate fields (the
        # registry treats _text/_specification/_remark consistently;
        # for this variant we use the _remark suffix mirroring
        # pii_q2_j's mandatory-remark pattern).
        bool_val = _three_state_boolean_from_row(
            cells, _PART_I_YES_COL, _PART_I_NO_COL,
        )
        fields_out[field_id] = GeneratedFieldValue(value=bool_val)
        fields_out[f"{field_id}_remark"] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_I_COMMENT_COL),
        )
    else:
        logger.warning(
            "gold_loader: unknown Part I input_type %r for field %r — "
            "extracting Yes/No only",
            input_type, field_id,
        )
        fields_out[field_id] = GeneratedFieldValue(
            value=_three_state_boolean_from_row(
                cells, _PART_I_YES_COL, _PART_I_NO_COL,
            ),
        )


def _process_part_ii_question(
    table, q_def: dict, fields_out: dict[str, GeneratedFieldValue],
    used_rows: set[int],
) -> None:
    """Process one Part II field. Part II has a Yes/No/N/A pattern."""
    field_id = q_def["field_id"]
    question = q_def.get("question", "")
    input_type = q_def.get("input_type", "")

    if not question:
        fields_out[field_id] = GeneratedFieldValue(value=None)
        return

    prefix = _first_n_words(question)
    row_idx = _find_row(table, prefix, used_rows)
    if row_idx is None:
        fields_out[field_id] = GeneratedFieldValue(value=None)
        return
    used_rows.add(row_idx)

    cells = table.rows[row_idx].cells

    if input_type == "yes_no_na":
        # N/A column present; if marked, value is None (per registry semantics)
        if _PART_II_NA_COL < len(cells) and _cell_is_marked(cells[_PART_II_NA_COL]):
            fields_out[field_id] = GeneratedFieldValue(value=None)
            return
        fields_out[field_id] = GeneratedFieldValue(
            value=_three_state_boolean_from_row(
                cells, _PART_II_YES_COL, _PART_II_NO_COL, _PART_II_NA_COL,
            ),
        )
    elif input_type == "yes_no_with_mandatory_remark":
        bool_val = _three_state_boolean_from_row(
            cells, _PART_II_YES_COL, _PART_II_NO_COL,
        )
        fields_out[field_id] = GeneratedFieldValue(value=bool_val)
        fields_out[f"{field_id}_remark"] = GeneratedFieldValue(
            value=_comment_value(cells, _PART_II_COMMENT_COL),
        )
    else:
        logger.warning(
            "gold_loader: unknown Part II input_type %r for field %r",
            input_type, field_id,
        )
        fields_out[field_id] = GeneratedFieldValue(
            value=_three_state_boolean_from_row(
                cells, _PART_II_YES_COL, _PART_II_NO_COL, _PART_II_NA_COL,
            ),
        )


def _load_part_i(table, manifest: dict) -> dict[str, GeneratedFieldValue]:
    fields: dict[str, GeneratedFieldValue] = {}
    used_rows: set[int] = set()

    for _q_key, q_def in manifest.get("part_i", {}).items():
        if not isinstance(q_def, dict):
            continue
        if "field_id" in q_def and "question" in q_def:
            _process_part_i_question(table, q_def, fields, used_rows)
        for sub in q_def.get("sub_fields", []) or []:
            if isinstance(sub, dict):
                _process_part_i_question(table, sub, fields, used_rows)
    return fields


def _load_part_ii(table, manifest: dict) -> dict[str, GeneratedFieldValue]:
    fields: dict[str, GeneratedFieldValue] = {}
    used_rows: set[int] = set()

    part_ii = manifest.get("part_ii", []) or []
    for f in part_ii:
        if isinstance(f, dict):
            _process_part_ii_question(table, f, fields, used_rows)
    return fields


def _load_signoff(table) -> dict[str, GeneratedFieldValue]:
    """Sign-off table layout (per renderer._render_signoff):
        R0: [engagement_partner_name, concurring_partner_name]
        R2 C0: sign_off_date
    """
    out: dict[str, GeneratedFieldValue] = {}
    try:
        ep = table.rows[0].cells[0].text.strip() or None
        cp = table.rows[0].cells[1].text.strip() or None
        sd = table.rows[2].cells[0].text.strip() or None
    except IndexError:
        ep, cp, sd = None, None, None
    out["engagement_partner"] = GeneratedFieldValue(value=ep)
    out["concurring_partner"] = GeneratedFieldValue(value=cp)
    out["sign_off_date"] = GeneratedFieldValue(value=sd)
    return out


# ---------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------

def load_filled_workpaper(
    filepath: str | Path,
    workpaper_type: str = "NPO-CX-1.1",
    engagement_id: str = "",
    manifest_path: str | Path | None = None,
) -> GeneratedWorkpaper:
    """Load a filled NPO-CX-1.1 .docx workpaper and produce a
    GeneratedWorkpaper.

    Parameters
    ----------
    filepath : str | Path
        Path to the filled .docx workpaper.
    workpaper_type : str
        Must be "NPO-CX-1.1" in Phase 2.2. Other workpaper types raise
        NotImplementedError until their loaders are added.
    engagement_id : str
        Optional engagement identifier to record on the result.
    manifest_path : str | Path | None
        Optional override for the manifest YAML path. Defaults to
        config/npo_cx_1_1_manifest.yaml.

    Returns
    -------
    GeneratedWorkpaper
        With all registry fields populated. Hybrid yes_no_with_X
        fields are split per the registry convention (boolean +
        _specification/_text/_reference/_remark). Citations are
        EMPTY — call citation_linker.auto_link_citations() to
        populate from matching ExtractedFact provenance.

    Raises
    ------
    FileNotFoundError, NotImplementedError
    """
    if workpaper_type != "NPO-CX-1.1":
        raise NotImplementedError(
            f"gold_loader: workpaper_type {workpaper_type!r} not supported "
            "in Phase 2.2 (NPO-CX-1.1 only). Sibling variants "
            "(GOV/FP/TRB) reuse this loader after manifest alignment."
        )

    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"gold_loader: {path} not found")

    manifest_p = Path(manifest_path) if manifest_path else _DEFAULT_MANIFEST_PATH
    with open(manifest_p) as f:
        manifest = yaml.safe_load(f)

    doc = Document(str(path))
    fields: dict[str, GeneratedFieldValue] = {}

    if len(doc.tables) > _T_HEADER:
        fields.update(_load_header(doc.tables[_T_HEADER]))
    if len(doc.tables) > _T_PART_I:
        fields.update(_load_part_i(doc.tables[_T_PART_I], manifest))
    if len(doc.tables) > _T_SIGNOFF:
        fields.update(_load_signoff(doc.tables[_T_SIGNOFF]))
    if len(doc.tables) > _T_PART_II:
        part_ii_fields = _load_part_ii(doc.tables[_T_PART_II], manifest)
        fields.update(part_ii_fields)

        # engagement_type is auditor_selection meta — not literally stored
        # in the .docx. We infer it from Part II population: if ANY Part II
        # field has a non-null value, this was an initial engagement
        # (Part II is filled for initial / 1st-year engagements only,
        # per the manifest's part_ii_behavior). Otherwise it's recurring.
        # The registry expects a categorical from the allowed_values list.
        has_part_ii_content = any(
            fv.value is not None for fv in part_ii_fields.values()
        )
        fields["engagement_type"] = GeneratedFieldValue(
            value=(
                "Initial / 1st Year"
                if has_part_ii_content
                else "Recurring / 2nd Year or Subsequent"
            ),
        )

    logger.info(
        "gold_loader: %s -> %d fields extracted (workpaper=%s, engagement=%s)",
        path.name, len(fields), workpaper_type, engagement_id,
    )

    return GeneratedWorkpaper(
        workpaper_type=workpaper_type,
        engagement_id=engagement_id,
        fields=fields,
    )
