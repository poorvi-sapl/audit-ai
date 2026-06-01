"""
raw_to_training_pair/generation/tests/test_phase_2_2.py
=========================================================
Phase 2.2 tests — locks in the real-data integration layer:
gold_loader, citation_linker, engagement_ingest, and PII enforcement
in pair_builder.

All tests use synthetic data — no real HCLLP engagement docs needed.
The gold_loader is exercised against a python-docx fixture built
in-test so we control every cell. The engagement_ingest is exercised
against a temp directory with synthetic files. PII enforcement uses
the real Presidio stack but with synthetic PII strings.

Run with:
    pytest raw_to_training_pair/generation/tests/test_phase_2_2.py -v
"""

from __future__ import annotations

import logging
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from auditai_data_normalization.assembly_layer import SourceDocumentExtraction
from auditai_data_normalization.generation_contract import (
    CHAR_OFFSET_UNAVAILABLE,
    ExtractedFact,
    SourceCitation,
)
from raw_to_training_pair.generation.citation_linker import (
    auto_link_citations,
)
from raw_to_training_pair.generation.engagement_ingest import (
    classify_document_type,
    ingest_engagement_folder,
)
from raw_to_training_pair.generation.gold_loader import (
    _has_mark,
    _normalize,
    _strip_label_prefix,
    load_filled_workpaper,
)
from raw_to_training_pair.generation.pair_builder import build_generation_pair
from raw_to_training_pair.generation.target_schema import (
    GeneratedCitation,
    GeneratedFieldValue,
    GeneratedWorkpaper,
)

WORKPAPER = "NPO-CX-1.1"


# ---------------------------------------------------------------------
# gold_loader — pure-function helpers
# ---------------------------------------------------------------------

class TestGoldLoaderHelpers:

    def test_normalize_lowercases_and_strips_punctuation(self):
        assert _normalize("Q1(a). Audit of FS?") == "q1 a audit of fs"

    def test_normalize_drops_hyphens(self):
        # "non-attest" → "nonattest" so it matches the manifest's normalized form
        assert _normalize("non-attest services") == "nonattest services"

    def test_strip_label_prefix_extracts_value(self):
        assert _strip_label_prefix(
            "Organization: Sample Foundation, Inc.", "Organization",
        ) == "Sample Foundation, Inc."

    def test_strip_label_prefix_returns_none_for_empty(self):
        assert _strip_label_prefix("Organization:", "Organization") is None

    def test_strip_label_prefix_handles_missing_label(self):
        # If the cell doesn't start with the label, return as-is
        assert _strip_label_prefix("Some other content", "Organization") == "Some other content"

    def test_has_mark_detects_X(self):
        assert _has_mark("X") is True
        assert _has_mark("  x  ") is True
        assert _has_mark("") is False
        assert _has_mark("Yes") is False


# ---------------------------------------------------------------------
# gold_loader — full-document round-trip with python-docx fixture
# ---------------------------------------------------------------------

def _build_synthetic_npo_docx(tmp_path: Path) -> Path:
    """Build a minimal NPO-CX-1.1-shaped .docx fixture in-test.

    Creates exactly the 5 expected tables with cell values that match
    what renderer.py would produce. Returns the path to the saved file.

    The loader doesn't need every Part I/II row populated — it'll skip
    rows it can't match by prefix. We populate just enough to verify
    the parsing logic on representative field types.
    """
    doc = Document()

    # Table 0: Header (2 rows x 2 cols — only the values we need)
    t_header = doc.add_table(rows=2, cols=2)
    t_header.rows[0].cells[0].text = "Organization: Sample Foundation, Inc."
    t_header.rows[0].cells[1].text = (
        "Statement of Financial Position Date: 2024-06-30"
    )
    t_header.rows[1].cells[0].text = "Completed by: Jane Auditor"
    t_header.rows[1].cells[1].text = "Date: 2024-09-15"

    # Table 1: Instructions placeholder (unused — but the loader expects
    # table 2 to be Part I, so this slot must exist)
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "(instructions)"

    # Table 2: Part I (a handful of rows we can match by prefix)
    t_part_i = doc.add_table(rows=8, cols=5)
    # Row 0: q1_c — "Single Audit?"
    t_part_i.rows[0].cells[0].text = "Single Audit?"
    t_part_i.rows[0].cells[2].text = "X"   # Yes column → true
    # Row 1: q1_a — "Audit of organization's financial statements..."
    t_part_i.rows[1].cells[0].text = (
        "Audit of organization's financial statements in accordance "
        "with GAAS? (Specify basis of accounting.)"
    )
    t_part_i.rows[1].cells[4].text = "Accrual / GAAP"  # dropdown value in Comments
    # Row 2: q3 — "Briefly describe the intended use..."
    t_part_i.rows[2].cells[0].text = (
        "Briefly describe the intended use of the financial statements."
    )
    t_part_i.rows[2].cells[4].text = (
        "Intended for use by management and stakeholders."
    )
    # Row 3: q4 — "Is the financial reporting framework... unacceptable?"
    t_part_i.rows[3].cells[0].text = (
        "Is the financial reporting framework used by management to "
        "prepare the financial statements considered unacceptable?"
    )
    t_part_i.rows[3].cells[3].text = "X"   # No → false
    # Row 4: q5_a — yes_no_with_text → boolean + _text
    t_part_i.rows[4].cells[0].text = (
        "Accept responsibility for the preparation and fair presentation "
        "of the financial statements or for the design, implementation, "
        "and maintenance of internal control over the financial statements?"
    )
    t_part_i.rows[4].cells[3].text = "X"   # No → false
    t_part_i.rows[4].cells[4].text = (
        "The management accepts responsibility..."
    )
    # Row 5: q1_g — yes_no_with_text
    t_part_i.rows[5].cells[0].text = "Other non-attest/non-audit services?"
    t_part_i.rows[5].cells[2].text = "X"
    t_part_i.rows[5].cells[4].text = (
        "Assist in preparing the financial statements"
    )
    # Row 6: q1_f — yes_no_with_specification
    t_part_i.rows[6].cells[0].text = (
        "Preparation of federal tax or information returns? (Specify.)"
    )
    t_part_i.rows[6].cells[2].text = "X"   # Yes
    t_part_i.rows[6].cells[4].text = "Form 990"   # specification dropdown
    # Row 7: q2 — yes_no_with_reference
    t_part_i.rows[7].cells[0].text = (
        "Will our firm provide non-attest/non-audit services to the client?"
    )
    t_part_i.rows[7].cells[2].text = "X"
    t_part_i.rows[7].cells[4].text = "NPO CX-1.2"

    # Table 3: Sign-off (4 rows x 2 cols)
    t_signoff = doc.add_table(rows=4, cols=2)
    t_signoff.rows[0].cells[0].text = "Sanwar Harshwal"   # engagement_partner
    t_signoff.rows[0].cells[1].text = "Jane Concurring"   # concurring_partner
    t_signoff.rows[2].cells[0].text = "2024-09-20"        # sign_off_date

    # Table 4: Part II (1 row exercising yes_no_na + 1 row for pii_q2_j)
    t_part_ii = doc.add_table(rows=2, cols=6)
    # pii_q3 — yes_no_na
    t_part_ii.rows[0].cells[0].text = (
        "Have contacts with bankers, attorneys, board members, funding "
        "sources, state attorney generals, credit services, or others "
        "having business relationships with the client raised any concerns "
        "about management's integrity or other concerns about the client?"
    )
    t_part_ii.rows[0].cells[3].text = "X"   # No
    # pii_q2_j — yes_no_with_mandatory_remark
    t_part_ii.rows[1].cells[0].text = (
        "Document (or cross-reference to documentation of) the "
        "predecessor's understanding as to the reasons for a change "
        "in auditors and any additional comments from the results of "
        "the inquiries of the predecessor auditor."
    )
    t_part_ii.rows[1].cells[2].text = "X"   # Yes
    t_part_ii.rows[1].cells[5].text = (
        "Predecessor confirmed reasons in writing on 2024-03-15."
    )

    out = tmp_path / "synthetic_filled_npo.docx"
    doc.save(out)
    return out


class TestGoldLoaderFullDoc:

    def test_loads_header_fields(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path, engagement_id="ENG-TEST")
        assert gw.engagement_id == "ENG-TEST"
        assert gw.fields["organization_name"].value == "Sample Foundation, Inc."
        assert gw.fields["financial_position_date"].value == "2024-06-30"
        assert gw.fields["completed_by"].value == "Jane Auditor"
        assert gw.fields["completion_date"].value == "2024-09-15"

    def test_loads_yes_no_field(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["q1_c"].value is True
        assert gw.fields["q4"].value is False

    def test_loads_dropdown_field(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["q1_a"].value == "Accrual / GAAP"

    def test_loads_text_prefill_field(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert "stakeholders" in (gw.fields["q3"].value or "")

    def test_splits_yes_no_with_text_into_two_fields(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        # q5_a → boolean (False from No mark) + _text
        assert gw.fields["q5_a"].value is False
        assert "accepts responsibility" in (gw.fields["q5_a_text"].value or "")
        # q1_g → boolean (True from Yes mark) + _text
        assert gw.fields["q1_g"].value is True
        assert "preparing the financial statements" in (gw.fields["q1_g_text"].value or "")

    def test_splits_yes_no_with_specification_into_two_fields(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["q1_f"].value is True
        assert gw.fields["q1_f_specification"].value == "Form 990"

    def test_splits_yes_no_with_reference_into_two_fields(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["q2"].value is True
        assert gw.fields["q2_reference"].value == "NPO CX-1.2"

    def test_loads_signoff_block(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["engagement_partner"].value == "Sanwar Harshwal"
        assert gw.fields["concurring_partner"].value == "Jane Concurring"
        assert gw.fields["sign_off_date"].value == "2024-09-20"

    def test_loads_part_ii_yes_no_na_field(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        # pii_q3 marked No → False
        assert gw.fields["pii_q3"].value is False

    def test_splits_yes_no_with_mandatory_remark(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        assert gw.fields["pii_q2_j"].value is True
        assert "2024-03-15" in (gw.fields["pii_q2_j_remark"].value or "")

    def test_no_citations_emitted_by_loader(self, tmp_path):
        # Decision L2: loader leaves citations empty; linker fills them
        path = _build_synthetic_npo_docx(tmp_path)
        gw = load_filled_workpaper(path)
        for fid, fv in gw.fields.items():
            assert fv.citations == [], f"loader populated citations for {fid}"

    def test_unsupported_workpaper_type_raises(self, tmp_path):
        path = _build_synthetic_npo_docx(tmp_path)
        with pytest.raises(NotImplementedError, match="GOV-CX-1.1"):
            load_filled_workpaper(path, workpaper_type="GOV-CX-1.1")

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            load_filled_workpaper(tmp_path / "nonexistent.docx")


# ---------------------------------------------------------------------
# citation_linker — auto-link from ExtractedFact provenance
# ---------------------------------------------------------------------

def _make_text_fact(field_id: str, value: str) -> ExtractedFact:
    return ExtractedFact(
        field_id=field_id,
        field_type="text",
        value=value,
        confidence=0.9,
        sources=[SourceCitation(
            document_path="/data/engagement_letter.pdf",
            document_type="engagement_letter",
            page=2,
            char_start=10, char_end=10 + len(value),
            quoted_text=value,
        )],
        extractor_method="structural_heuristic",
    )


class TestCitationLinker:

    def test_links_when_values_match(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={
                "organization_name": GeneratedFieldValue(value="Sample Foundation, Inc."),
            },
        )
        facts = {
            "organization_name": _make_text_fact(
                "organization_name", "Sample Foundation, Inc.",
            ),
        }
        linked = auto_link_citations(gold, facts)
        cites = linked.fields["organization_name"].citations
        assert len(cites) == 1
        assert cites[0].document == "engagement_letter"
        assert cites[0].page == 2

    def test_skips_when_values_differ(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={
                "organization_name": GeneratedFieldValue(value="Different Org, LLC"),
            },
        )
        facts = {
            "organization_name": _make_text_fact(
                "organization_name", "Sample Foundation, Inc.",
            ),
        }
        linked = auto_link_citations(gold, facts)
        assert linked.fields["organization_name"].citations == []

    def test_skips_null_gold_values(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={"organization_name": GeneratedFieldValue(value=None)},
        )
        facts = {
            "organization_name": _make_text_fact(
                "organization_name", "Anything",
            ),
        }
        linked = auto_link_citations(gold, facts)
        assert linked.fields["organization_name"].citations == []

    def test_skips_booleans_by_default(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={"q1_c": GeneratedFieldValue(value=False)},
        )
        # Create a synthetic boolean fact (need ExtractedFact for boolean type)
        bool_fact = ExtractedFact(
            field_id="q1_c", field_type="boolean", value=False, confidence=0.9,
            sources=[SourceCitation(
                document_path="/x.pdf", document_type="engagement_letter",
                page=1, char_start=0, char_end=2, quoted_text="No",
            )],
            extractor_method="structural_heuristic",
        )
        linked = auto_link_citations(gold, {"q1_c": bool_fact})
        # Default: booleans skipped, no citation copied
        assert linked.fields["q1_c"].citations == []

    def test_links_booleans_when_opted_in(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={"q1_c": GeneratedFieldValue(value=False)},
        )
        bool_fact = ExtractedFact(
            field_id="q1_c", field_type="boolean", value=False, confidence=0.9,
            sources=[SourceCitation(
                document_path="/x.pdf", document_type="engagement_letter",
                page=1, char_start=0, char_end=2, quoted_text="No",
            )],
            extractor_method="structural_heuristic",
        )
        linked = auto_link_citations(
            gold, {"q1_c": bool_fact}, copy_boolean_citations=True,
        )
        assert len(linked.fields["q1_c"].citations) == 1

    def test_does_not_mutate_original_gold(self):
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={
                "organization_name": GeneratedFieldValue(value="Sample Foundation, Inc."),
            },
        )
        facts = {
            "organization_name": _make_text_fact(
                "organization_name", "Sample Foundation, Inc.",
            ),
        }
        auto_link_citations(gold, facts)
        # Original gold's citations stay empty
        assert gold.fields["organization_name"].citations == []

    def test_preserves_pre_existing_citations(self):
        existing = GeneratedCitation(document="manual", page=1, quoted_text="x")
        gold = GeneratedWorkpaper(
            workpaper_type=WORKPAPER, engagement_id="ENG-X",
            fields={
                "organization_name": GeneratedFieldValue(
                    value="Sample Foundation, Inc.",
                    citations=[existing],
                ),
            },
        )
        facts = {
            "organization_name": _make_text_fact(
                "organization_name", "Sample Foundation, Inc.",
            ),
        }
        linked = auto_link_citations(gold, facts)
        # Pre-existing + newly-linked are both present
        cites = linked.fields["organization_name"].citations
        assert any(c.document == "manual" for c in cites)
        assert any(c.document == "engagement_letter" for c in cites)


# ---------------------------------------------------------------------
# engagement_ingest — classification + folder routing
# ---------------------------------------------------------------------

class TestEngagementIngestClassifier:

    @pytest.mark.parametrize("name,expected", [
        ("engagement_letter_2024.pdf",  "engagement_letter"),
        ("Engagement-Letter.pdf",        "engagement_letter"),
        ("prior_year_audit.pdf",         "prior_year_file"),
        ("client_intake_form.pdf",       "client_intake_form"),
        ("financial_statements_2024.pdf","financial_statements"),
        ("audit_report.pdf",             "audit_report"),
        ("board_minutes_2024-03.pdf",    "board_minutes"),
        ("random_file.pdf",              "unknown"),
    ])
    def test_classify_document_type(self, name, expected):
        assert classify_document_type(name) == expected


class TestEngagementIngestFolder:

    def test_folder_not_found_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            ingest_engagement_folder(tmp_path / "nope")

    def test_not_a_directory_raises(self, tmp_path):
        f = tmp_path / "file.txt"
        f.write_text("x")
        with pytest.raises(NotADirectoryError):
            ingest_engagement_folder(f)

    def test_empty_folder_returns_empty_list(self, tmp_path):
        result = ingest_engagement_folder(tmp_path)
        assert result == []

    def test_skips_unsupported_extensions(self, tmp_path):
        (tmp_path / "engagement_letter.docx").write_bytes(b"fake docx")
        (tmp_path / "notes.txt").write_text("ignored")
        (tmp_path / ".hidden").write_text("ignored")
        # Patch docx_extractor.extract so we don't actually parse the fake bytes
        with patch(
            "auditai_data_normalization.extractors.docx_extractor.extract"
        ) as mock_ex:
            fake_record = MagicMock()
            fake_record.pages_text = None
            mock_ex.return_value = fake_record
            extractions = ingest_engagement_folder(tmp_path)
        # Only the .docx is ingested; .txt and .hidden are skipped
        assert len(extractions) == 1
        assert extractions[0].document_type == "engagement_letter"

    def test_doc_type_hint_overrides_filename_classification(self, tmp_path):
        (tmp_path / "random_name.docx").write_bytes(b"fake")
        with patch(
            "auditai_data_normalization.extractors.docx_extractor.extract"
        ) as mock_ex:
            fake_record = MagicMock()
            fake_record.pages_text = None
            mock_ex.return_value = fake_record
            extractions = ingest_engagement_folder(
                tmp_path,
                doc_type_hints={"random_name.docx": "engagement_letter"},
            )
        assert extractions[0].document_type == "engagement_letter"

    def test_file_extractor_failure_is_logged_and_skipped(self, tmp_path, caplog):
        (tmp_path / "engagement_letter.docx").write_bytes(b"fake")
        with patch(
            "auditai_data_normalization.extractors.docx_extractor.extract",
            side_effect=RuntimeError("parsing failed"),
        ):
            with caplog.at_level(logging.WARNING):
                extractions = ingest_engagement_folder(tmp_path)
        assert extractions == []
        assert any("parsing failed" in r.message for r in caplog.records)


# ---------------------------------------------------------------------
# pair_builder — PII enforcement (Decision Y2)
# ---------------------------------------------------------------------

def _minimal_gen_input():
    from raw_to_training_pair.generation.tests.test_generation_pipeline import (
        _synthetic_gen_input,
    )
    return _synthetic_gen_input()


def _minimal_gold():
    from raw_to_training_pair.generation.tests.test_generation_pipeline import (
        _synthetic_gold_minimal_npo,
    )
    return _synthetic_gold_minimal_npo()


class TestPairBuilderPIIEnforcement:

    @patch("raw_to_training_pair.generation.pair_builder._check_pii_in_content")
    def test_pii_issues_recorded_in_metadata_in_warn_mode(self, mock_check):
        mock_check.return_value = (3, ["PERSON", "DATE_TIME"])
        pair = build_generation_pair(
            _minimal_gen_input(), _minimal_gold(), pii_strict=False,
        )
        assert pair["metadata"]["pii_issues"]["user_pii_count"] == 3
        assert "PERSON" in pair["metadata"]["pii_issues"]["user_pii_types"]

    @patch("raw_to_training_pair.generation.pair_builder._check_pii_in_content")
    def test_pii_strict_raises_when_pii_detected(self, mock_check):
        mock_check.return_value = (1, ["PERSON"])
        with pytest.raises(ValueError, match="PII detected"):
            build_generation_pair(
                _minimal_gen_input(), _minimal_gold(), pii_strict=True,
            )

    @patch("raw_to_training_pair.generation.pair_builder._check_pii_in_content")
    def test_pii_strict_allows_when_no_pii(self, mock_check):
        mock_check.return_value = (0, [])
        pair = build_generation_pair(
            _minimal_gen_input(), _minimal_gold(), pii_strict=True,
        )
        assert pair["metadata"]["pii_issues"]["user_pii_count"] == 0
        assert pair["metadata"]["pii_issues"]["assistant_pii_count"] == 0

    @patch("raw_to_training_pair.generation.pair_builder._check_pii_in_content")
    def test_pii_warn_mode_logs_warning(self, mock_check, caplog):
        mock_check.return_value = (2, ["PERSON"])
        with caplog.at_level(logging.WARNING):
            build_generation_pair(
                _minimal_gen_input(), _minimal_gold(), pii_strict=False,
            )
        assert any("PII detected" in r.message for r in caplog.records)

    def test_metadata_carries_pii_issues_key_when_clean(self):
        # No PII in fully-synthetic test content; metadata should still have
        # the pii_issues key (with zero counts) for downstream consumers.
        pair = build_generation_pair(_minimal_gen_input(), _minimal_gold())
        assert "pii_issues" in pair["metadata"]
        # In the synthetic test fixtures the gold contains "Sample Foundation, Inc."
        # and "Jane Auditor" — Presidio MAY flag these as PERSON entities.
        # The test only asserts the structure exists, not the counts.
        assert "user_pii_count" in pair["metadata"]["pii_issues"]
        assert "assistant_pii_count" in pair["metadata"]["pii_issues"]
