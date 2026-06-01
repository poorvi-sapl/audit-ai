"""
raw_to_training_pair/generation/tests/test_phase_2_3.py
=========================================================
Phase 2.3 tests — batch orchestrator (walks engagement folders,
produces generation pairs, appends to JSONL) and task balance
reporter (audits a JSONL corpus for multi-task balance).

All tests use synthetic data. The batch orchestrator's full
pipeline (ingest → load → link → assemble → build → append) is
exercised against in-test python-docx fixtures and mocked file
extractors. The task balance reporter is tested against synthetic
JSONL files.

Run with:
    pytest raw_to_training_pair/generation/tests/test_phase_2_3.py -v
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from raw_to_training_pair.generation.batch_orchestrator import (
    BatchResult,
    _find_gold_workpaper,
    run_batch_from_folder,
)
from raw_to_training_pair.generation.task_balance_reporter import (
    TaskBalanceReport,
    pretty_print,
    report_task_balance,
)


# ---------------------------------------------------------------------
# _find_gold_workpaper
# ---------------------------------------------------------------------

class TestFindGoldWorkpaper:

    def test_returns_none_when_no_match(self, tmp_path):
        (tmp_path / "other.pdf").write_bytes(b"x")
        assert _find_gold_workpaper(tmp_path, "*filled*.docx") is None

    def test_returns_single_match(self, tmp_path):
        gold = tmp_path / "filled_workpaper.docx"
        gold.write_bytes(b"fake")
        assert _find_gold_workpaper(tmp_path, "*filled*.docx") == gold

    def test_raises_on_ambiguous(self, tmp_path):
        (tmp_path / "filled_a.docx").write_bytes(b"a")
        (tmp_path / "filled_b.docx").write_bytes(b"b")
        with pytest.raises(ValueError, match="multiple files matched"):
            _find_gold_workpaper(tmp_path, "*filled*.docx")


# ---------------------------------------------------------------------
# Synthetic NPO-CX-1.1 .docx fixture (shared with test_phase_2_2)
# ---------------------------------------------------------------------

def _build_synthetic_npo_docx_minimal(tmp_path: Path, filename: str) -> Path:
    """Minimal synthetic filled NPO-CX-1.1. Just enough tables for the
    loader to find header + sign-off. We don't need every Part I row
    populated for the orchestrator smoke test — gold_loader gracefully
    emits null GeneratedFieldValue for unmatched rows.
    """
    doc = Document()
    # Table 0: Header
    t_h = doc.add_table(rows=2, cols=2)
    t_h.rows[0].cells[0].text = "Organization: Sample Foundation, Inc."
    t_h.rows[0].cells[1].text = "Statement of Financial Position Date: 2024-06-30"
    t_h.rows[1].cells[0].text = "Completed by: Jane Auditor"
    t_h.rows[1].cells[1].text = "Date: 2024-09-15"
    # Table 1: Instructions placeholder
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "(instructions)"
    # Table 2: Part I (empty rows; loader will emit nulls)
    doc.add_table(rows=2, cols=5)
    # Table 3: Sign-off
    t_s = doc.add_table(rows=4, cols=2)
    t_s.rows[0].cells[0].text = "Sanwar Harshwal"
    t_s.rows[0].cells[1].text = ""
    t_s.rows[2].cells[0].text = "2024-09-20"
    # Table 4: Part II
    doc.add_table(rows=1, cols=6)
    out = tmp_path / filename
    doc.save(out)
    return out


def _make_engagement_folder(
    root: Path, eng_id: str, with_gold: bool = True,
    gold_filename: str = "filled_npo.docx",
) -> Path:
    """Create an engagement subdirectory with a synthetic gold + maybe
    a source doc."""
    folder = root / eng_id
    folder.mkdir()
    # A non-gold "source" file so the ingest layer has something to scan
    (folder / "engagement_letter.txt").write_text("synthetic source")
    if with_gold:
        _build_synthetic_npo_docx_minimal(folder, gold_filename)
    return folder


# ---------------------------------------------------------------------
# run_batch_from_folder — happy path + edge cases
# ---------------------------------------------------------------------

class TestRunBatchFromFolder:

    def test_root_not_found_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            run_batch_from_folder(
                tmp_path / "nonexistent",
                tmp_path / "out.jsonl",
            )

    def test_root_not_a_directory_raises(self, tmp_path):
        f = tmp_path / "file.txt"
        f.write_text("x")
        with pytest.raises(NotADirectoryError):
            run_batch_from_folder(f, tmp_path / "out.jsonl")

    def test_empty_root_returns_empty_result(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        result = run_batch_from_folder(root, tmp_path / "out.jsonl")
        assert result.total_engagements == 0
        assert result.pairs_built == 0

    def test_skips_dot_and_underscore_subfolders(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        (root / ".scratch").mkdir()
        (root / "_archive").mkdir()
        _make_engagement_folder(root, "ENG-001")
        result = run_batch_from_folder(
            root, tmp_path / "out.jsonl",
            gold_filename_pattern="*filled*.docx",
        )
        assert result.total_engagements == 1

    def test_missing_gold_recorded_as_error(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        _make_engagement_folder(root, "ENG-001", with_gold=False)
        result = run_batch_from_folder(
            root, tmp_path / "out.jsonl",
            gold_filename_pattern="*filled*.docx",
        )
        assert result.total_engagements == 1
        assert result.pairs_built == 0
        assert any("no gold" in e.lower() for e in result.errors)
        assert result.per_engagement[0].success is False

    def test_single_engagement_end_to_end_builds_and_writes(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        _make_engagement_folder(root, "ENG-001")
        out = tmp_path / "out.jsonl"
        result = run_batch_from_folder(
            root, out, gold_filename_pattern="*filled*.docx",
        )
        assert result.total_engagements == 1
        assert result.pairs_built == 1
        assert result.pairs_written == 1
        assert out.exists()
        # JSONL has exactly one line with valid JSON
        lines = [ln for ln in out.read_text().splitlines() if ln.strip()]
        assert len(lines) == 1
        pair = json.loads(lines[0])
        assert pair["metadata"]["pair_type"] == "generation"
        assert pair["metadata"]["engagement_id"] == "ENG-001"

    def test_two_engagements_two_pairs(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        _make_engagement_folder(root, "ENG-001")
        _make_engagement_folder(root, "ENG-002")
        out = tmp_path / "out.jsonl"
        result = run_batch_from_folder(root, out, gold_filename_pattern="*filled*.docx")
        assert result.total_engagements == 2
        assert result.pairs_built == 2
        assert result.pairs_written == 2

    def test_idempotent_rerun_deduplicates(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        _make_engagement_folder(root, "ENG-001")
        out = tmp_path / "out.jsonl"
        # First run writes
        r1 = run_batch_from_folder(root, out, gold_filename_pattern="*filled*.docx")
        assert r1.pairs_written == 1
        # Second run should dedup (same engagement → same pair_hash)
        r2 = run_batch_from_folder(root, out, gold_filename_pattern="*filled*.docx")
        assert r2.pairs_built == 1
        assert r2.pairs_written == 0
        assert r2.pairs_deduplicated == 1

    def test_per_engagement_failure_does_not_block_others(self, tmp_path):
        root = tmp_path / "engagements"
        root.mkdir()
        _make_engagement_folder(root, "ENG-001", with_gold=False)  # will fail
        _make_engagement_folder(root, "ENG-002", with_gold=True)   # will succeed
        result = run_batch_from_folder(
            root, tmp_path / "out.jsonl",
            gold_filename_pattern="*filled*.docx",
        )
        assert result.total_engagements == 2
        assert result.pairs_built == 1
        assert len(result.errors) == 1


# ---------------------------------------------------------------------
# task_balance_reporter
# ---------------------------------------------------------------------

def _write_jsonl_pairs(
    path: Path, pair_specs: list[tuple[str, int]],
) -> None:
    """Write a JSONL with `count` pairs of each pair_type. Each pair has
    a unique pair_hash so they don't dedup if appended."""
    lines: list[str] = []
    counter = 0
    for pair_type, count in pair_specs:
        for _ in range(count):
            counter += 1
            pair = {
                "messages": [
                    {"role": "system", "content": "x"},
                    {"role": "user", "content": "y"},
                    {"role": "assistant", "content": "z"},
                ],
                "metadata": {
                    "pair_type": pair_type,
                    "task": (
                        "GENERATE_WORKPAPER" if pair_type == "generation"
                        else "REVIEW_WORKPAPER"
                    ),
                    "pair_hash": f"hash_{counter:06d}",
                },
            }
            lines.append(json.dumps(pair))
    path.write_text("\n".join(lines) + "\n")


class TestReportTaskBalance:

    def test_file_not_found_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            report_task_balance(tmp_path / "nope.jsonl")

    def test_empty_file_returns_zero_total(self, tmp_path):
        p = tmp_path / "empty.jsonl"
        p.write_text("")
        r = report_task_balance(p)
        assert r.total_pairs == 0
        assert r.counts == {}

    def test_blank_lines_skipped(self, tmp_path):
        p = tmp_path / "blanks.jsonl"
        p.write_text("\n\n\n")
        r = report_task_balance(p)
        assert r.total_pairs == 0

    def test_malformed_line_skipped_with_warning(self, tmp_path):
        p = tmp_path / "bad.jsonl"
        p.write_text("not valid json\n")
        r = report_task_balance(p)
        assert r.total_pairs == 0
        assert any("malformed" in w.lower() for w in r.warnings)

    def test_balanced_50_50_two_types(self, tmp_path):
        p = tmp_path / "balanced.jsonl"
        _write_jsonl_pairs(p, [("review", 5), ("generation", 5)])
        r = report_task_balance(p, max_share=0.40)
        assert r.total_pairs == 10
        assert r.counts == {"review": 5, "generation": 5}
        assert abs(r.shares["review"] - 0.5) < 1e-6
        # 50% > 40% threshold → warning for BOTH
        assert len(r.warnings) == 2

    def test_balanced_under_threshold_no_warnings(self, tmp_path):
        p = tmp_path / "ok.jsonl"
        _write_jsonl_pairs(p, [
            ("review", 3),       # 30%
            ("generation", 4),   # 40% — exactly at threshold (not over)
            ("validation", 3),   # 30%
        ])
        r = report_task_balance(p, max_share=0.40)
        assert r.total_pairs == 10
        # 0.40 is NOT > 0.40 — exactly at threshold passes
        assert r.is_balanced

    def test_imbalanced_triggers_warning(self, tmp_path):
        p = tmp_path / "imbal.jsonl"
        _write_jsonl_pairs(p, [
            ("review", 8),       # 80% — way over 40%
            ("generation", 2),   # 20%
        ])
        r = report_task_balance(p, max_share=0.40)
        assert not r.is_balanced
        assert any("review" in w for w in r.warnings)
        assert all("generation" not in w for w in r.warnings)

    def test_taxonomy_task_buckets_differently(self, tmp_path):
        p = tmp_path / "task.jsonl"
        _write_jsonl_pairs(p, [("review", 3), ("generation", 7)])
        r_pair = report_task_balance(p, taxonomy="pair_type")
        r_task = report_task_balance(p, taxonomy="task")
        assert "review" in r_pair.counts
        assert "REVIEW_WORKPAPER" in r_task.counts

    def test_missing_taxonomy_key_bucketed_as_missing(self, tmp_path):
        p = tmp_path / "no_pair_type.jsonl"
        pair = {
            "messages": [],
            "metadata": {"pair_hash": "abc"},  # no pair_type
        }
        p.write_text(json.dumps(pair) + "\n")
        r = report_task_balance(p)
        assert "(missing)" in r.counts

    def test_pretty_print_includes_total_and_distribution(self, tmp_path):
        p = tmp_path / "ok.jsonl"
        _write_jsonl_pairs(p, [("review", 3), ("generation", 7)])
        r = report_task_balance(p, max_share=0.40)
        out = pretty_print(r)
        assert "Total pairs: 10" in out
        assert "review" in out
        assert "generation" in out

    def test_pretty_print_flags_violations(self, tmp_path):
        p = tmp_path / "imbal.jsonl"
        _write_jsonl_pairs(p, [("review", 9), ("generation", 1)])
        r = report_task_balance(p, max_share=0.40)
        out = pretty_print(r)
        assert "⚠️" in out
        assert "Warnings" in out
