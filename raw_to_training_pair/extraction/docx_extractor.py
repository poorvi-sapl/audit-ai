"""
raw_to_training_pair/extraction/docx_extractor.py
===================================================
Phase 2 raw extractor for .docx files.

Contract
--------
Returns a plain dict — no DocumentRecord, no PII scrubbing, no confidence
scoring. This is raw material for the pair builder only.

Output shape
------------
{
    "text"    : str,   # full document text, paragraphs joined by newline
    "tables"  : list,  # [{"headers": [...], "rows": [[...], ...]}, ...]
    "metadata": dict   # heading_count, table_count, paragraph_count
}

Public API
----------
    extract(file_path: str | Path) -> dict
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Heading style names to detect section headings
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3", "heading 4",
    "heading 5", "heading 6", "h1", "h2", "h3", "title", "subtitle",
}


def _get_cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _extract_tables(doc: Document) -> list[dict]:
    """Extract all tables as list of {headers, rows} dicts."""
    tables = []
    for tbl in doc.tables:
        raw_rows = []
        for row in tbl.rows:
            cells = [_get_cell_text(cell) for cell in row.cells]
            # Forward-fill merged cells
            filled = []
            prev = ""
            for i, cell in enumerate(cells):
                if cell == prev and i > 0:
                    filled.append("")
                else:
                    filled.append(cell)
                    prev = cell
            if any(c.strip() for c in filled):
                raw_rows.append(filled)

        if not raw_rows:
            continue

        headers = raw_rows[0]
        rows = raw_rows[1:] if len(raw_rows) > 1 else []
        tables.append({"headers": headers, "rows": rows})

    return tables


def extract(file_path: str | Path) -> dict:
    """
    Extract raw text and tables from a .docx file.

    Parameters
    ----------
    file_path : str | Path

    Returns
    -------
    dict with keys: text, tables, metadata

    Raises
    ------
    FileNotFoundError
        If file does not exist.
    RuntimeError
        If python-docx cannot open the file.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise RuntimeError(f"Failed to open {path.name}: {e}") from e

    # --- Extract paragraphs ---
    paragraphs = []
    heading_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name in _HEADING_STYLES:
            heading_count += 1
            paragraphs.append(f"\n## {text}")
        else:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    # --- Extract tables ---
    tables = _extract_tables(doc)

    # Append table text to full_text so pair builder can read it
    for i, tbl in enumerate(tables):
        lines = [" | ".join(tbl["headers"])]
        for row in tbl["rows"]:
            lines.append(" | ".join(row))
        full_text += f"\n\n[TABLE {i + 1}]\n" + "\n".join(lines)

    metadata = {
        "heading_count": heading_count,
        "table_count": len(tables),
        "paragraph_count": len(paragraphs),
    }

    logger.info(
        "p2/docx_extractor: %s — %d paragraphs, %d tables",
        path.name, len(paragraphs), len(tables),
    )

    return {"text": full_text, "tables": tables, "metadata": metadata}