"""
workpaper_generator/renderer.py
=================================
Writes resolved field values into a copy of the NPO-CX-1.1 .docx template
and saves the filled workpaper alongside the engagement's trace.json.

Table layout (from the existing NPO-CX-1.1 templates):
  Table 0  Header        (3 rows x 2 cols) — org name, FYE date, completed by, date
  Table 1  Instructions  (1 row  x 1 col)  — narrative, untouched
  Table 2  Part I        (152 rows x 5 cols) — Q1–Q16, cols: [text][_][Yes][No][Comments]
  Table 3  Sign-off      (4 rows x 2 cols) — engagement partner, dates
  Table 4  Part II       (23 rows x 6 cols) — Initial-engagement Qs, cols: [text][_][Yes][No][N/A][Comments]

Matching strategy:
  Manifest field questions are matched to table rows by normalized
  first-N-word prefix. Rows containing "practical consideration" or
  "practical considerations" are skipped (they are SOP notes, not
  questions).
"""

from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path
from typing import Optional

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from workpaper_generator.orchestrator import EngagementResult
from workpaper_generator.rule_engine import ResolvedField

_PROJECT_ROOT = Path(__file__).parent.parent
_DEFAULT_MANIFEST_PATH = _PROJECT_ROOT / "config" / "npo_cx_1_1_manifest.yaml"

# Mark text written into Yes / No / N/A cells.
_MARK = "X"

# Table indices in the .docx
_T_HEADER = 0
_T_PART_I = 2
_T_SIGNOFF = 3
_T_PART_II = 4

# Column indices
_PART_I_QUESTION_COL = 0
_PART_I_YES_COL = 2
_PART_I_NO_COL = 3
_PART_I_COMMENT_COL = 4

_PART_II_QUESTION_COL = 0
_PART_II_YES_COL = 2
_PART_II_NO_COL = 3
_PART_II_NA_COL = 4
_PART_II_COMMENT_COL = 5

_SKIP_ROW_MARKERS = ("practical consideration",)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _strip_annotation_runs(doc: Document, question_cols: list[tuple[int, int]]) -> int:
    """
    Remove runs with explicit non-auto font color from question-text cells.
    These are team annotation runs (blue 0000FF etc.) added to the blank
    template — they should never appear in the final generated workpaper.
    question_cols: list of (table_index, col_index) to target.
    Returns count of runs removed.
    """
    removed = 0
    for t_idx, col_idx in question_cols:
        if t_idx >= len(doc.tables):
            continue
        for row in doc.tables[t_idx].rows:
            cell = row.cells[col_idx]
            for p in cell.paragraphs:
                to_remove = []
                for run in p.runs:
                    color_el = run._element.find(f".//{{{_W}}}color")
                    if color_el is not None:
                        val = color_el.get(f"{{{_W}}}val", "auto")
                        if val.lower() != "auto":
                            to_remove.append(run._element)
                for el in to_remove:
                    el.getparent().remove(el)
                    removed += 1
    return removed


# XML styleId values (w:pStyle w:val) that render as visible bullets in the
# NPO-CX-1.1 blank template but carry no content.
_BULLET_STYLE_IDS = {"CXContent", "CXGutter", "CXStepContent"}


def _strip_empty_bullets(doc: Document) -> int:
    """
    Neutralise empty bullet/list paragraphs in all table cells.
    Word requires at least one w:p per cell so these cannot be deleted.
    Instead their style is reset to Normal (w:pStyle removed) so the
    bullet character stops rendering while the required cell structure
    is preserved.
    Returns the number of paragraphs neutralised.
    """
    _p = f"{{{_W}}}p"
    _t = f"{{{_W}}}t"
    _pPr = f"{{{_W}}}pPr"
    _pStyle = f"{{{_W}}}pStyle"
    _numPr = f"{{{_W}}}numPr"
    _val = f"{{{_W}}}val"

    neutralised = 0
    for tc in doc.element.iter(f"{{{_W}}}tc"):
        for p_el in tc.findall(_p):
            text = "".join(t.text or "" for t in p_el.iter(_t)).strip()
            if text:
                continue
            has_num = p_el.find(f".//{_numPr}") is not None
            pPr_el = p_el.find(_pPr)
            style_id = ""
            if pPr_el is not None:
                ps = pPr_el.find(_pStyle)
                if ps is not None:
                    style_id = ps.get(_val, "")
            if not (has_num or style_id in _BULLET_STYLE_IDS):
                continue
            # Reset to Normal: remove pStyle and any numPr
            if pPr_el is not None:
                for ps in pPr_el.findall(_pStyle):
                    pPr_el.remove(ps)
                for nm in pPr_el.findall(_numPr):
                    pPr_el.remove(nm)
                # Remove pPr entirely if now empty
                if len(pPr_el) == 0 and not pPr_el.text:
                    p_el.remove(pPr_el)
            neutralised += 1
    return neutralised


def _normalize(text: str) -> str:
    text = text.lower()
    # Drop hyphens entirely so 'non-attest' matches 'nonattest'.
    text = text.replace("-", "")
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _first_n_words(text: str, n: int = 8) -> str:
    return " ".join(_normalize(text).split()[:n])


def _row_text(row) -> str:
    return row.cells[0].text


def _is_skippable_row(row_text_norm: str) -> bool:
    return any(m in row_text_norm for m in _SKIP_ROW_MARKERS) or not row_text_norm


def _set_cell_valign(cell: _Cell, val: str) -> None:
    """Set vertical alignment on a table cell (top | center | bottom)."""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.find(f"{{{_W}}}tcPr")
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    vAlign = tcPr.find(f"{{{_W}}}vAlign")
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(f"{{{_W}}}val", val)


def _set_cell_text(cell: _Cell, text: str, space_after_pt: int = 0) -> None:
    """Replace cell contents with text, preserving the first paragraph style.

    `space_after_pt` adds breathing room below the paragraph so that
    adjacent short rows (e.g. Q1(a) → Q1(b)) don't visually collapse
    into each other.
    """
    # Wipe existing content
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph(text)
    if space_after_pt:
        p.paragraph_format.space_after = Pt(space_after_pt)


def _append_cell_text(cell: _Cell, text: str) -> None:
    """Append a paragraph of text to a cell, leaving existing content intact."""
    cell.add_paragraph(text)


def _clear_cell(cell: _Cell) -> None:
    """Wipe a cell entirely, leaving one empty paragraph for layout."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    cell.add_paragraph("")


def _drop_appended_paragraphs(cell: _Cell) -> None:
    """Drop every paragraph after the first.

    Used on the question-text column (col 0) of matched rows to strip
    any auditor-added answers that may have been embedded as additional
    paragraphs by previous edits to the source template.
    """
    paras = list(cell.paragraphs)
    if len(paras) > 1:
        for p in paras[1:]:
            p._element.getparent().remove(p._element)


def _build_question_index(manifest: dict) -> dict[str, str]:
    """Map field_id → normalized question prefix used for table matching."""
    index: dict[str, str] = {}

    for f in manifest.get("header_fields", []):
        index[f["field_id"]] = ""  # header has its own table layout

    part_i = manifest.get("part_i", {})
    for q_key, q_def in part_i.items():
        if "field_id" in q_def and "question" in q_def:
            index[q_def["field_id"]] = _first_n_words(q_def["question"])
        for sub in q_def.get("sub_fields", []) if isinstance(q_def, dict) else []:
            if "question" in sub:
                index[sub["field_id"]] = _first_n_words(sub["question"])

    for f in manifest.get("part_ii", []):
        if "question" in f:
            index[f["field_id"]] = _first_n_words(f["question"])

    return index


def _find_row(table, target_prefix: str, used_rows: set[int]) -> Optional[int]:
    """
    Find the first row whose first-8-word normalized prefix starts with
    `target_prefix`. Skips rows already matched to another field and
    skips SOP "Practical Consideration" notes.
    """
    if not target_prefix:
        return None
    for ri, row in enumerate(table.rows):
        if ri in used_rows:
            continue
        c0_norm = _normalize(_row_text(row))
        if _is_skippable_row(c0_norm):
            continue
        c0_prefix = _first_n_words(_row_text(row))
        if c0_prefix.startswith(target_prefix):
            return ri
    return None


def _render_header(table, resolved: dict[str, ResolvedField]) -> None:
    """
    Header table layout:
      R0 C0: 'Organization: ...'
      R0 C1: 'Statement of Financial Position Date: ...'
      R1 C0: 'Completed by: ...'
      R1 C1: 'Date: ...'
    We rewrite the prefix-style content with the resolved value (or leave
    the placeholder if needs_input).
    """
    mapping = [
        (0, 0, "Organization", resolved.get("organization_name")),
        (0, 1, "Statement of Financial Position Date", resolved.get("financial_position_date")),
        (1, 0, "Completed by", resolved.get("completed_by")),
        (1, 1, "Date", resolved.get("completion_date")),
    ]
    for row_idx, col_idx, label, rf in mapping:
        cell = table.rows[row_idx].cells[col_idx]
        if rf and rf.status == "resolved" and rf.value:
            _set_cell_text(cell, f"{label}: {rf.value}")
        # Else leave the existing label/placeholder untouched.


def _mark_part_i_row(row, resolved_value: str, comment: Optional[str] = None) -> None:
    cells = row.cells
    if resolved_value in ("Yes", "No"):
        target_col = _PART_I_YES_COL if resolved_value == "Yes" else _PART_I_NO_COL
        _set_cell_text(cells[target_col], _MARK)
    if comment:
        _set_cell_text(cells[_PART_I_COMMENT_COL], comment)


def _mark_part_ii_row(row, resolved_value: str, comment: Optional[str] = None) -> None:
    cells = row.cells
    if resolved_value == "Yes":
        _set_cell_text(cells[_PART_II_YES_COL], _MARK)
    elif resolved_value == "No":
        _set_cell_text(cells[_PART_II_NO_COL], _MARK)
    elif resolved_value == "N/A":
        _set_cell_text(cells[_PART_II_NA_COL], _MARK)
    if comment:
        _set_cell_text(cells[_PART_II_COMMENT_COL], comment)


def _comment_for_field(rf: ResolvedField) -> Optional[str]:
    """Build the Comments-column text for a resolved field."""
    if rf.status == "na":
        return None

    val = rf.value or ""

    if rf.source == "sop_fixed":
        if " — " in val:
            return val.split(" — ", 1)[1]
        if val not in ("Yes", "No", "N/A"):
            return val

    if rf.source == "sop_fixed_plus_lookup":
        return val or None

    if rf.source == "sop_fixed_plus_manual":
        return val.replace("Yes — ", "").strip() or None

    if rf.source == "auditor_selection":
        # "Yes — Form 990"             → "Form 990" (Q1f spec)
        # "No"                          → no comment
        # "Accrual / GAAP"              → "Accrual / GAAP" (Q1a basis)
        # "2 CFR Part 200 – Uniform..." → "2 CFR Part 200 – Uniform..." (Q1b)
        if val.startswith("Yes — "):
            return val[len("Yes — "):].strip() or None
        if val in ("Yes", "No", ""):
            return None
        return val

    if rf.source == "auditor_yes_no_with_remark":
        # "Yes — <remark>" → comment is the remark (Q12)
        # "No"             → no comment
        if val.startswith("Yes — "):
            return val[len("Yes — "):].strip() or None
        return None

    return None


def _value_for_marking(rf: ResolvedField) -> Optional[str]:
    """Return 'Yes' / 'No' / 'N/A' for cell marking, ignoring extras."""
    if rf.status == "na":
        return "N/A"
    if not rf.value:
        return None
    if rf.value.startswith("Yes"):
        return "Yes"
    if rf.value.startswith("No"):
        return "No"
    if rf.value == "N/A":
        return "N/A"
    # An auditor_selection value that doesn't start with Yes/No is a pure
    # specification (Q1a basis of accounting, Q1b grant framework) — the
    # question is answered affirmatively with that specification.
    if rf.source == "auditor_selection":
        return "Yes"
    return None


def _render_part_i(table, resolved: dict[str, ResolvedField], question_index: dict[str, str], manifest: dict) -> dict[str, str]:
    """Fill Part I question rows. Returns {field_id: 'matched_row' | 'not_found'} for diagnostics."""
    placements: dict[str, str] = {}
    used_rows: set[int] = set()

    part_i_field_ids = []
    for q_key, q_def in manifest["part_i"].items():
        if "field_id" in q_def:
            part_i_field_ids.append(q_def["field_id"])
        for sub in q_def.get("sub_fields", []) if isinstance(q_def, dict) else []:
            part_i_field_ids.append(sub["field_id"])

    for fid in part_i_field_ids:
        rf = resolved.get(fid)
        if not rf:
            placements[fid] = "no_resolved"
            continue
        prefix = question_index.get(fid, "")
        ri = _find_row(table, prefix, used_rows)
        if ri is None:
            placements[fid] = "not_found"
            continue
        used_rows.add(ri)

        # Defensive cleanup: wipe any residue from the source template
        # (Yes/No marks, comments, auditor answers appended to col 0)
        # before placing our marks. Keeps output deterministic.
        row_cells = table.rows[ri].cells
        _drop_appended_paragraphs(row_cells[_PART_I_QUESTION_COL])
        _clear_cell(row_cells[_PART_I_YES_COL])
        _clear_cell(row_cells[_PART_I_NO_COL])
        _clear_cell(row_cells[_PART_I_COMMENT_COL])

        mark = _value_for_marking(rf)
        comment = _comment_for_field(rf)
        # Q3 has no Yes/No — pure prefill text → write to Comments only.
        # Both cells get vAlign=top so the long comment and short question
        # start at the same height instead of the question appearing below.
        if fid == "q3":
            if rf.value:
                _set_cell_text(row_cells[_PART_I_COMMENT_COL], rf.value, space_after_pt=6)
            _set_cell_valign(row_cells[_PART_I_QUESTION_COL], "top")
            _set_cell_valign(row_cells[_PART_I_COMMENT_COL], "top")
            placements[fid] = f"row_{ri}_text"
            continue
        if comment:
            _set_cell_text(row_cells[_PART_I_COMMENT_COL], comment, space_after_pt=6)

        if mark:
            _mark_part_i_row(table.rows[ri], mark, None)
            placements[fid] = f"row_{ri}_{mark}{'_with_comment' if comment else ''}"
        elif comment:
            placements[fid] = f"row_{ri}_comment_only"
        else:
            placements[fid] = f"row_{ri}_left_blank"

    return placements


def _render_part_ii(table, resolved: dict[str, ResolvedField], question_index: dict[str, str], manifest: dict) -> dict[str, str]:
    placements: dict[str, str] = {}
    used_rows: set[int] = set()
    for f in manifest["part_ii"]:
        fid = f["field_id"]
        rf = resolved.get(fid)
        if not rf:
            placements[fid] = "no_resolved"
            continue
        prefix = question_index.get(fid, "")
        ri = _find_row(table, prefix, used_rows)
        if ri is None:
            placements[fid] = "not_found"
            continue
        used_rows.add(ri)

        # Defensive cleanup before writing
        row_cells = table.rows[ri].cells
        _drop_appended_paragraphs(row_cells[_PART_II_QUESTION_COL])
        _clear_cell(row_cells[_PART_II_YES_COL])
        _clear_cell(row_cells[_PART_II_NO_COL])
        _clear_cell(row_cells[_PART_II_NA_COL])
        _clear_cell(row_cells[_PART_II_COMMENT_COL])

        mark = _value_for_marking(rf)
        comment = _comment_for_field(rf)
        if comment:
            _set_cell_text(row_cells[_PART_II_COMMENT_COL], comment, space_after_pt=6)

        if mark:
            _mark_part_ii_row(table.rows[ri], mark, None)
            placements[fid] = f"row_{ri}_{mark}{'_with_comment' if comment else ''}"
        elif comment:
            placements[fid] = f"row_{ri}_comment_only"
        else:
            placements[fid] = f"row_{ri}_left_blank"
    return placements


def _render_acceptance_decision_paragraph(doc, resolved: dict[str, ResolvedField]) -> bool:
    """
    Append a bolded line right after the SOP paragraph
        "We should accept/continue   or not accept/continue   the engagement."
    surfacing the system's computed acceptance decision and the rule that fired.

    The original SOP paragraph is preserved untouched so the audit trail and
    SOP wording stay intact; the new line is a clearly visible computed result.

    Returns True if the line was appended, False if the target paragraph
    couldn't be located or the decision isn't resolved.
    """
    decision_rf = resolved.get("acceptance_decision")
    if not decision_rf or not decision_rf.value:
        return False

    target = None
    for p in doc.paragraphs:
        text = p.text.lower()
        if "should accept/continue" in text and "not accept/continue" in text:
            target = p
            break
    if target is None:
        return False

    new_p_xml = OxmlElement("w:p")
    target._element.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, target._parent)

    bold_run = new_p.add_run(f"System decision: {decision_rf.value}")
    bold_run.bold = True

    if decision_rf.rule_applied:
        rule_run = new_p.add_run(f"  ({decision_rf.rule_applied})")
        rule_run.italic = True

    return True


def _render_signoff(table, resolved: dict[str, ResolvedField]) -> None:
    """
    Sign-off table layout:
      R0: ['Sanwar Harshwal', '']                                ← engagement partner / concurring partner names
      R1: ['Engagement Partner', 'Concurring Partner (if required)']
      R2: ['', '']                                               ← date row 1
      R3: ['Date', 'Date']
    The template already has 'Sanwar Harshwal' in R0 C0. We fill in sign_off_date in R2 C0 if resolved.
    """
    sign_off = resolved.get("sign_off_date")
    if sign_off and sign_off.status == "resolved" and sign_off.value:
        _set_cell_text(table.rows[2].cells[0], sign_off.value)

    concurring = resolved.get("concurring_partner")
    if concurring and concurring.status == "resolved" and concurring.value:
        _set_cell_text(table.rows[0].cells[1], concurring.value)


def _load_manifest(path: str | Path) -> dict:
    with open(path) as f:
        return yaml.safe_load(f)


def render(
    result: EngagementResult,
    manifest_path: str | Path = _DEFAULT_MANIFEST_PATH,
    output_filename: Optional[str] = None,
) -> Path:
    """
    Render the resolved fields into a copy of the engagement's template.
    Returns the path to the written .docx.
    """
    manifest = _load_manifest(manifest_path)
    question_index = _build_question_index(manifest)

    template_src = Path(result.workpaper_template_path)
    output_name = output_filename or f"{result.engagement_name}_NPO_CX_1.1_filled.docx"
    output_path = Path(result.output_dir) / output_name
    shutil.copy(template_src, output_path)

    doc = Document(output_path)
    _strip_annotation_runs(doc, [
        (_T_PART_I,  _PART_I_QUESTION_COL),
        (_T_PART_II, _PART_II_QUESTION_COL),
    ])

    _render_header(doc.tables[_T_HEADER], result.resolved)
    part_i_placements = _render_part_i(doc.tables[_T_PART_I], result.resolved, question_index, manifest)
    _render_signoff(doc.tables[_T_SIGNOFF], result.resolved)
    part_ii_placements = _render_part_ii(doc.tables[_T_PART_II], result.resolved, question_index, manifest)
    decision_written = _render_acceptance_decision_paragraph(doc, result.resolved)
    bullets_removed = _strip_empty_bullets(doc)

    doc.save(output_path)

    # Write a placement diagnostic next to the .docx
    diag_path = Path(result.output_dir) / "placement_diagnostic.txt"
    with open(diag_path, "w") as f:
        f.write("Part I placements:\n")
        for fid, where in part_i_placements.items():
            f.write(f"  {fid}: {where}\n")
        f.write("\nPart II placements:\n")
        for fid, where in part_ii_placements.items():
            f.write(f"  {fid}: {where}\n")
        f.write(f"\nAcceptance decision paragraph written: {decision_written}\n")
        f.write(f"Empty bullet paragraphs removed: {bullets_removed}\n")

    return output_path
