"""
streamlit_app/app.py
=====================
AuditAI Training Pair Generator — Streamlit UI.

Single-page interface for:
0. SOP Management   — upload SOP documents, embed into Qdrant
1. Upload & Process — upload a workpaper, extract fields, generate pairs
2. Review Queue     — approve / reject pairs across all sessions
3. Export           — write approved pairs to JSONL, download files

UI calls pipeline functions only. Zero business logic here.

Run with:
    streamlit run streamlit_app/app.py
"""

import sys
import tempfile
from pathlib import Path

import streamlit as st

sys.path.insert(0, str(Path(__file__).parent.parent))

from pipeline.pipeline import process_workpaper, write_approved
from raw_to_training_pair.auditor_review import (
    load_pending,
    approve,
    conditional_approve,
    send_for_correction,
    reject,
    stats,
)
from raw_to_training_pair.jsonl_writer import count, get_output_path
from auditai_data_normalization.alias_suggester import (
    load_suggestions,
    approve as alias_approve,
    reject as alias_reject,
    coverage_stats,
    suggest_canonical,
    _load_canonical_fields,
)
from auditai_data_normalization.normalize import reset_alias_cache
from workpaper_generator.orchestrator import (
    run_engagement,
    register_new_client,
    ENGAGEMENT_INITIAL,
    ENGAGEMENT_RECURRING,
    CANONICAL_BLANK_TEMPLATE,
)
from workpaper_generator.renderer import render as render_workpaper
from workpaper_generator.pdf_section_detector import detect as detect_sections
from workpaper_generator.rule_engine import resolve_workpaper

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_QUEUE_PATH = Path("data/review_queue.jsonl")
_DATA_DIR   = Path("data")
_CLIENT_TYPES = ["NPO", "Government", "For-Profit", "Tribal"]
_SOP_EXTENSIONS = ["pdf", "txt", "md", "docx"]
_ENGAGEMENTS_DIR = Path("Engagement Accept and Cont Form")
_WORKPAPER_RUNS_DIR = Path("data/workpaper_runs")

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="AuditAI — Training Pair Generator",
    page_icon="📋",
    layout="wide",
)

# ---------------------------------------------------------------------------
# Session state initialisation
# ---------------------------------------------------------------------------

defaults = {
    "process_result": None,
    "reviewer_id": "",
    "last_export": None,
    "queue_refresh": 0,
    "sop_embed_result": None,
    "workpaper_result": None,
    "workpaper_output_path": None,
    "playground_detection": None,
    "playground_resolved": None,
    "playground_pdf_label": None,
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _review_score_stats(queue_path: Path) -> dict:
    """
    Read review_confidence from every pair in the queue and return
    aggregated stats split by status.

    Returns dict with keys: all, approved, pending — each a list of floats.
    Returns empty lists if queue file doesn't exist or has no scores.
    """
    import json

    buckets: dict[str, list[float]] = {"all": [], "approved": [], "pending": []}
    if not queue_path.exists():
        return buckets

    with open(queue_path, encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if not line:
                continue
            try:
                entry = json.loads(line)
            except json.JSONDecodeError:
                continue
            # Queue format: {"pair": {..., "metadata": {...}}, "status": "pending", ...}
            pair_data = entry.get("pair", entry)
            score = pair_data.get("metadata", {}).get("review_confidence")
            if score is None:
                continue
            score = float(score)
            status = entry.get("status", "pending")
            buckets["all"].append(score)
            if status in ("approved", "conditional"):
                buckets["approved"].append(score)
            elif status == "pending":
                buckets["pending"].append(score)

    return buckets


def _qdrant_chunk_count(collection_name: str) -> int:
    """
    Return number of vectors in the Qdrant collection.
    Returns:
        >= 0  : connected, value is chunk count (0 if collection not yet created)
        -1    : Qdrant unreachable (connection error)
    """
    try:
        import os
        from qdrant_client import QdrantClient
        client = QdrantClient(
            host=os.getenv("QDRANT_HOST", "localhost"),
            port=int(os.getenv("QDRANT_PORT", "6333")),
            timeout=5,
        )
        # Confirm connection by listing collections first
        existing = {c.name for c in client.get_collections().collections}
        if collection_name not in existing:
            return 0   # Connected but collection not yet created — not an error
        info = client.get_collection(collection_name)
        return getattr(info, "vectors_count", None) or getattr(info, "points_count", 0) or 0
    except Exception:
        return -1


def _extract_text_from_sop(file_path: Path) -> str:
    """Extract raw text from SOP file for chunking."""
    suffix = file_path.suffix.lower()
    if suffix in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages.append(text)
            return "\n\n".join(pages)
        except Exception as e:
            st.error(f"PDF extraction failed: {e}")
            return ""
    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            parts = []

            # Body paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # Tables — bug #15: paragraphs-only missed all table content
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [
                        " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                        for cell in row.cells
                    ]
                    row_text = "\t".join(c for c in row_cells if c)
                    if row_text.strip():
                        parts.append(row_text)

            # Headers and footers from each section
            for section in doc.sections:
                for hf in [section.header, section.footer]:
                    if hf:
                        for p in hf.paragraphs:
                            if p.text.strip():
                                parts.append(p.text.strip())

            return "\n".join(parts)
        except Exception as e:
            st.error(f"DOCX extraction failed: {e}")
            return ""
    return ""


# ---------------------------------------------------------------------------
# Sidebar
# ---------------------------------------------------------------------------

with st.sidebar:
    st.title("⚙️ Settings")

    use_mock = st.checkbox(
        "Use mock completions",
        value=False,
        help=(
            "ON  → semi-deterministic mock (fast, no Ollama needed)\n"
            "OFF → real Gemma 3 12B via Ollama (requires ollama serve)"
        ),
    )

    if use_mock:
        st.info("Mock mode — completions are simulated")
    else:
        st.warning("Real mode — Ollama must be running")

    st.divider()

    # Qdrant SOP chunk count — wrapped so a missing Qdrant never blocks page load
    try:
        from config.settings import get_settings
        settings = get_settings()
        qdrant_count = _qdrant_chunk_count(settings.qdrant.collection_sop)
        if qdrant_count >= 0:
            st.metric("SOP chunks in Qdrant", qdrant_count)
        else:
            st.caption("⚠️ Qdrant unreachable")
    except Exception:
        st.caption("⚠️ Qdrant unreachable")

    st.divider()
    st.subheader("📊 Queue stats")

    queue_stats = stats(_QUEUE_PATH)
    col1, col2 = st.columns(2)
    col1.metric("Total",       queue_stats["total"])
    col1.metric("Pending",     queue_stats["pending"])
    col2.metric("Approved",    queue_stats["approved"]
                               + queue_stats.get("conditional", 0))
    col2.metric("Needs rework",queue_stats.get("correction", 0))
    col1.metric("Rejected",    queue_stats["rejected"])

    if st.button("Show avg review scores", key="score_stats_btn", use_container_width=True):
        _scores = _review_score_stats(_QUEUE_PATH)

        def _fmt(vals: list[float]) -> str:
            if not vals:
                return "—"
            return f"{sum(vals) / len(vals):.3f}"

        s_col1, s_col2, s_col3 = st.columns(3)
        s_col1.metric("All",      _fmt(_scores["all"]),      help=f"{len(_scores['all'])} pairs")
        s_col2.metric("Approved", _fmt(_scores["approved"]), help=f"{len(_scores['approved'])} pairs")
        s_col3.metric("Pending",  _fmt(_scores["pending"]),  help=f"{len(_scores['pending'])} pairs")

        if _scores["all"]:
            import pandas as pd
            _df = pd.DataFrame({"score": _scores["all"]})
            _bins = [0.0, 0.5, 0.7, 0.85, 1.01]
            _labels = ["0–0.5", "0.5–0.7", "0.7–0.85", "0.85–1.0"]
            _df["bucket"] = pd.cut(_df["score"], bins=_bins, labels=_labels, right=False)
            _dist = _df["bucket"].value_counts().reindex(_labels, fill_value=0)
            st.bar_chart(_dist, height=120)
        else:
            st.caption("No scored pairs in queue yet.")

    st.divider()
    st.subheader("🏷️ Alias coverage")

    alias_csv = _DATA_DIR / "suggested_aliases.csv"
    alias_stats = coverage_stats(alias_csv)

    if alias_stats.total_seen == 0:
        st.caption("No unknown labels seen yet.")
    else:
        st.metric(
            "Coverage",
            f"{alias_stats.coverage_pct:.1f}%",
            help=f"{alias_stats.approved} of {alias_stats.total_seen} unknown labels mapped",
        )
        acol1, acol2 = st.columns(2)
        acol1.metric(
            "Pending", alias_stats.pending,
            delta=f"{alias_stats.high_conf} high-conf" if alias_stats.high_conf else None,
        )
        acol2.metric("Approved", alias_stats.approved)

    st.divider()
    st.caption("AuditAI — HCLLP Internal Tool")

# ---------------------------------------------------------------------------
# Main header
# ---------------------------------------------------------------------------

st.title("📋 AuditAI Training Pair Generator")
st.caption("Manage SOPs → Upload workpapers → Generate pairs → Review → Export JSONL")

st.divider()

# ===========================================================================
# SECTION 0 — SOP Management
# ===========================================================================

st.header("0 · SOP Management")
st.caption("Upload SOP documents to embed into Qdrant. Do this once before processing workpapers.")

sop_col1, sop_col2 = st.columns([2, 1])

with sop_col1:
    uploaded_sop = st.file_uploader(
        "Upload a SOP document",
        type=_SOP_EXTENSIONS,
        help="Supported: .pdf, .txt, .md, .docx",
        key="sop_uploader",
    )

with sop_col2:
    sop_version = st.text_input(
        "SOP version",
        value="2024-Q1",
        placeholder="e.g. 2024-Q1, v3.2",
        help="Version tag for this SOP — used for versioning in Qdrant and Postgres",
    )

embed_btn = st.button(
    "🔵 Embed SOP into Qdrant",
    disabled=(uploaded_sop is None or not sop_version.strip()),
    type="primary",
    key="embed_sop_btn",
)

if embed_btn and uploaded_sop and sop_version.strip():
    _DATA_DIR.mkdir(parents=True, exist_ok=True)
    tmp_sop_path = _DATA_DIR / uploaded_sop.name

    with open(tmp_sop_path, "wb") as f:
        f.write(uploaded_sop.getbuffer())

    with st.spinner(f"Embedding {uploaded_sop.name} — loading model and chunking..."):
        try:
            from engineering_benchmark.sop_chunker import chunk_text
            from engineering_benchmark.embedder import embed_chunks

            # Extract text
            sop_text = _extract_text_from_sop(tmp_sop_path)

            if not sop_text.strip():
                st.error("Could not extract text from this SOP file.")
            else:
                # Chunk
                chunks = chunk_text(
                    text=sop_text,
                    source_doc=uploaded_sop.name,
                    sop_version=sop_version.strip(),
                )

                # Embed and upsert
                embed_result = embed_chunks(chunks, pg_conn=None)
                st.session_state.sop_embed_result = {
                    "file_name":  uploaded_sop.name,
                    "sop_version": sop_version.strip(),
                    "chunks_created": len(chunks),
                    "upserted": embed_result.upserted,
                    "skipped": embed_result.skipped,
                    "errors": embed_result.errors,
                }

        except ImportError as e:
            st.error(
                f"sentence-transformers not installed: {e}\n"
                "Run: pip install -e '.[training]'"
            )
        except Exception as e:
            st.error(f"Embedding failed: {e}")

# Show SOP embed result
if st.session_state.sop_embed_result is not None:
    res = st.session_state.sop_embed_result
    if res["errors"]:
        st.warning(
            f"⚠️ Embedded with errors — "
            f"chunks: {res['chunks_created']} | "
            f"upserted: {res['upserted']} | "
            f"errors: {len(res['errors'])}"
        )
        with st.expander("Embedding errors"):
            for err in res["errors"]:
                st.error(err)
    else:
        st.success(
            f"✅ **{res['file_name']}** embedded successfully — "
            f"{res['chunks_created']} chunks created, "
            f"{res['upserted']} upserted to Qdrant "
            f"(version: {res['sop_version']})"
        )

st.divider()

# ===========================================================================
# SECTION 0.5 — Alias Suggestions
# ===========================================================================

st.header("0.5 · Alias Suggestions")
st.caption(
    "Unknown field labels found during extraction. "
    "Approve to add them to field_aliases.yaml. "
    "Reject to stop suggesting. Takes ~10 min per batch."
)

_alias_csv = _DATA_DIR / "suggested_aliases.csv"
_alias_pending = [
    r for r in load_suggestions(_alias_csv)
    if r.get("status") == "pending"
]

if not _alias_pending:
    st.info("✅ No pending alias suggestions — all labels are mapped.")
else:
    # ── Batch approve high-confidence ─────────────────────────────────────
    _high_conf = [
        s for s in _alias_pending
        if s.get("llm_confidence") == "high" and s.get("suggested_canonical")
    ]
    if _high_conf:
        st.markdown(f"**{len(_high_conf)} high-confidence suggestions** ready for batch approval")
        if st.button(
            f"✅ Batch approve {len(_high_conf)} high-confidence suggestions",
            disabled=not st.session_state.reviewer_id,
            key="batch_approve_btn",
        ):
            _n = sum(
                1 for s in _high_conf
                if alias_approve(s["raw_label"], s["suggested_canonical"],
                                 st.session_state.reviewer_id, _alias_csv)
            )
            reset_alias_cache()
            st.success(f"Approved {_n} aliases — field_aliases.yaml updated.")
            st.rerun()
        if not st.session_state.reviewer_id:
            st.caption("⚠️ Enter your reviewer ID in Section 2 to approve")
        st.divider()

    # ── Per-row review table ───────────────────────────────────────────────
    st.markdown(f"**{len(_alias_pending)} pending suggestions**")
    _canonical_fields = _load_canonical_fields()

    for _i, _s in enumerate(_alias_pending):
        _raw        = _s.get("raw_label", "")
        _sample     = _s.get("extracted_value", "")
        _src        = _s.get("source_file", "")
        _suggested  = _s.get("suggested_canonical", "")
        _llm_conf   = _s.get("llm_confidence", "pending")
        _seen       = _s.get("seen_count", "1")

        _conf_icon = {"high": "🟢", "medium": "🟡", "low": "🔴",
                      "none": "⚫", "pending": "⏳"}.get(_llm_conf, "⏳")

        with st.expander(
            f"{_conf_icon} `{_raw}` — seen {_seen}× "
            f"| suggestion: `{_suggested or 'none'}` ({_llm_conf})",
            expanded=(_llm_conf == "high"),
        ):
            _dc1, _dc2 = st.columns([2, 2])
            with _dc1:
                st.markdown(f"**Raw label:** `{_raw}`")
                st.markdown(f"**Sample value:** `{_sample or '—'}`")
                st.markdown(f"**Source:** `{_src or '—'}`")
                st.markdown(f"**Seen:** {_seen}×")

            with _dc2:
                _sel_key = f"alias_select_{_i}_{_raw[:10]}"
                _def_idx = 0
                if _suggested and _suggested in _canonical_fields:
                    try:
                        _def_idx = _canonical_fields.index(_suggested)
                    except ValueError:
                        _def_idx = 0

                _selected = st.selectbox(
                    "Map to canonical field",
                    options=[""] + _canonical_fields,
                    index=_def_idx + 1 if _suggested else 0,
                    key=_sel_key,
                )

                if not _suggested and st.button(
                    "🤖 Ask Gemma",
                    key=f"llm_btn_{_i}_{_raw[:8]}",
                    disabled=not st.session_state.reviewer_id,
                ):
                    with st.spinner("Asking Gemma..."):
                        _res = suggest_canonical(_raw, _canonical_fields)
                    if _res.has_suggestion:
                        st.info(
                            f"Gemma suggests: `{_res.suggested_canonical}` "
                            f"(confidence: {_res.llm_confidence})\n\n"
                            f"Reason: {_res.reasoning}"
                        )
                    else:
                        st.warning("Gemma found no confident match.")

            _bc1, _bc2, _ = st.columns([1, 1, 4])
            with _bc1:
                if st.button(
                    "✅ Approve",
                    key=f"alias_approve_{_i}_{_raw[:8]}",
                    disabled=(not st.session_state.reviewer_id or not _selected),
                ):
                    if alias_approve(_raw, _selected, st.session_state.reviewer_id, _alias_csv):
                        reset_alias_cache()
                        st.success(f"Mapped `{_raw}` → `{_selected}`")
                        st.rerun()
                    else:
                        st.error("Approval failed.")
            with _bc2:
                if st.button(
                    "❌ Reject",
                    key=f"alias_reject_{_i}_{_raw[:8]}",
                    disabled=not st.session_state.reviewer_id,
                ):
                    if alias_reject(_raw, st.session_state.reviewer_id, _alias_csv):
                        st.warning(f"Rejected `{_raw}`")
                        st.rerun()
                    else:
                        st.error("Rejection failed.")

st.divider()

# ===========================================================================
# SECTION 1 — Upload & Process
# ===========================================================================

st.header("1 · Upload & Process")

col_upload, col_options = st.columns([2, 1])

with col_upload:
    uploaded_file = st.file_uploader(
        "Upload a workpaper",
        type=["docx", "xlsx", "xls", "pdf", "csv", "json"],
        help="Supported: .docx, .xlsx, .xls, .pdf, .csv, .json",
    )

with col_options:
    st.markdown("**Client types to generate**")
    selected_client_types = []
    for ct in _CLIENT_TYPES:
        if st.checkbox(ct, value=(ct == "NPO"), key=f"ct_{ct}"):
            selected_client_types.append(ct)

    if not selected_client_types:
        st.warning("Select at least one client type")

process_btn = st.button(
    "⚡ Process workpaper",
    disabled=(uploaded_file is None or not selected_client_types),
    type="primary",
)

if process_btn and uploaded_file and selected_client_types:
    _DATA_DIR.mkdir(parents=True, exist_ok=True)
    tmp_path = _DATA_DIR / uploaded_file.name

    with open(tmp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    with st.spinner(f"Processing {uploaded_file.name}..."):
        result = process_workpaper(
            file_path=str(tmp_path),
            queue_path=_QUEUE_PATH,
            data_dir=_DATA_DIR,
            run_parallel=False,
            client_types=selected_client_types,
            use_mock=use_mock,
        )

    st.session_state.process_result = result
    st.session_state.queue_refresh += 1

# Show results
if st.session_state.process_result is not None:
    result = st.session_state.process_result

    if result.skipped:
        st.error(f"⚠️ Skipped: {result.skip_reason}")
    else:
        st.success(
            f"✅ Processed **{result.file_name}** — "
            f"{result.pairs_queued} pairs queued, "
            f"{result.pairs_failed} failed"
        )

        st.subheader("Extraction summary")
        record = result.record

        ecol1, ecol2, ecol3 = st.columns(3)
        ecol1.metric(
            "Confidence",
            f"{record.extraction_confidence:.2f}",
            delta="✓ pass" if record.extraction_confidence >= 0.7 else "✗ low",
        )
        ecol2.metric("Word count", record.word_count)
        ecol3.metric("File type", record.file_type)

        conf_summary = record.metadata.get("confidence_summary", {})
        fields_present = conf_summary.get("fields_present", [])
        fields_missing = conf_summary.get("fields_missing", [])
        pii_redactions = record.pii_redactions

        fcol1, fcol2 = st.columns(2)

        with fcol1:
            st.markdown(f"**Fields found ({len(fields_present)})**")
            if fields_present:
                st.code("\n".join(fields_present))
            else:
                st.caption("None detected")

        with fcol2:
            st.markdown(f"**Fields missing ({len(fields_missing)})**")
            if fields_missing:
                st.code("\n".join(fields_missing))
            else:
                st.caption("None — all fields found ✓")

        if pii_redactions:
            st.markdown("**PII redacted**")
            pii_summary = ", ".join(
                f"{r.pii_type}: {r.count}" for r in pii_redactions
            )
            st.info(f"🔒 {pii_summary}")

        if result.gate_failures:
            with st.expander(f"⚠️ Gate failures ({len(result.gate_failures)})"):
                for gf in result.gate_failures:
                    st.warning(gf)

st.divider()

# ===========================================================================
# SECTION 2 — Review Queue
# ===========================================================================

st.header("2 · Review Queue")

reviewer_id = st.text_input(
    "Your reviewer ID",
    value=st.session_state.reviewer_id,
    placeholder="e.g. SH, MS1, JD",
    help="Required to approve or reject pairs",
    key="reviewer_id_input",
)
st.session_state.reviewer_id = reviewer_id

pending_entries = load_pending(_QUEUE_PATH)

if not pending_entries:
    st.info("No pairs pending review.")
else:
    st.markdown(f"**{len(pending_entries)} pairs awaiting review**")

    for i, entry in enumerate(pending_entries):
        pair = entry.get("pair", {})
        meta = pair.get("metadata", {})
        messages = pair.get("messages", [])

        pair_hash = meta.get("pair_hash", "")
        label = (
            f"{meta.get('file_name', 'unknown')} · "
            f"{meta.get('client_type', '')} · "
            f"{meta.get('pair_type', '')} · "
            f"stage={meta.get('stage', '')} · "
            f"conf={meta.get('extraction_confidence', 0):.2f}"
        )

        with st.expander(f"#{i + 1} — {label}"):

            # ── E2 — Uncertainty summary banner ───────────────────────────
            uncertain       = meta.get("uncertain_sections", []) or []
            flagged_fields  = meta.get("flagged_fields", []) or []
            fields_missing  = meta.get("fields_missing", []) or []
            rev_conf        = meta.get("review_confidence", 0.0)
            llm_asst        = meta.get("llm_assisted", False)
            struct_evidence = meta.get("structural_evidence", {}) or {}

            # Collect all fields needing attention for the banner
            _all_attention = sorted(set(
                list(uncertain) + list(flagged_fields) + list(fields_missing)
            ))

            if _all_attention:
                st.warning(
                    f"⚠️ **{len(_all_attention)} field(s) need attention:** "
                    + ", ".join(f"`{f}`" for f in _all_attention)
                )
            if llm_asst:
                st.info("🤖 LLM-assisted extraction — verify flagged fields before approving")

            # Review confidence + extraction confidence row
            _rc_col, _ec_col = st.columns(2)
            if rev_conf > 0:
                _rc_col.caption(
                    f"Review confidence: **{rev_conf:.2f}** "
                    f"{'✓' if rev_conf >= 0.70 else '⚠️ below threshold'}"
                )
            _ext_conf = meta.get("extraction_confidence", 0.0)
            if _ext_conf > 0:
                _ec_col.caption(
                    f"Extraction confidence: **{_ext_conf:.2f}** "
                    f"{'✓' if _ext_conf >= 0.50 else '⚠️ low'}"
                )

            # ── E2 — Flagged fields detail panel ──────────────────────────
            if flagged_fields or fields_missing or struct_evidence:
                with st.expander(
                    f"🔍 Field detail — "
                    f"{len(flagged_fields)} flagged · "
                    f"{len(fields_missing)} missing · "
                    f"{len(struct_evidence)} structural"
                ):
                    if fields_missing:
                        st.markdown("**Missing fields** *(not found by any extractor)*")
                        for _f in fields_missing:
                            st.markdown(f"- ❌ `{_f}`")

                    if flagged_fields:
                        st.markdown("**Flagged fields** *(low confidence or LLM-only)*")
                        for _f in flagged_fields:
                            _tag = "🤖" if llm_asst else "⚠️"
                            st.markdown(f"- {_tag} `{_f}`")

                    if struct_evidence:
                        st.markdown("**Structural evidence** *(Phase 4 heuristic extraction)*")
                        for _fname, _ev in struct_evidence.items():
                            st.markdown(
                                f"- `{_fname}`: **{_ev.get('value', '')}** "
                                f"— conf={_ev.get('confidence', 0):.2f} "
                                f"p{_ev.get('source_page', '?')} "
                                f"via `{_ev.get('method', '?')}`"
                            )

            tab_sys, tab_user, tab_asst = st.tabs(["System", "User", "Assistant"])

            with tab_sys:
                sys_content = next(
                    (m["content"] for m in messages if m["role"] == "system"), ""
                )
                st.text_area("System prompt", sys_content, height=150,
                             disabled=True, key=f"sys_{i}")

            with tab_user:
                user_content = next(
                    (m["content"] for m in messages if m["role"] == "user"), ""
                )
                st.text_area("User message", user_content, height=250,
                             disabled=True, key=f"user_{i}")

            with tab_asst:
                asst_content = next(
                    (m["content"] for m in messages if m["role"] == "assistant"), ""
                )

                # ── E2 — Highlight placeholder SOP citations ──────────────
                import re as _re
                _placeholder_re = _re.compile(r"§X\.X", _re.IGNORECASE)
                _placeholder_count = len(_placeholder_re.findall(asst_content))

                if _placeholder_count > 0:
                    st.warning(
                        f"⚠️ **{_placeholder_count} placeholder SOP citation(s)** — "
                        "findings marked `§X.X` need real SOP sections before approving"
                    )
                    _annotated = _placeholder_re.sub("**[⚠️ §X.X — needs citation]**", asst_content)
                    st.markdown("**Completion preview** *(placeholder citations highlighted)*")
                    st.markdown(
                        f"<div style='background:#fff8e1;border-left:3px solid #f9a825;"
                        f"padding:10px;font-family:monospace;font-size:0.85em;"
                        f"white-space:pre-wrap;'>{_annotated}</div>",
                        unsafe_allow_html=True,
                    )
                    st.divider()

                # ── E2 — Inline edit (always available) ───────────────────
                edited_key = f"edit_{i}_{pair_hash[:8]}"
                if edited_key not in st.session_state:
                    st.session_state[edited_key] = asst_content
                edited_completion = st.text_area(
                    "Assistant completion (editable — changes apply on approve)",
                    value=st.session_state[edited_key],
                    height=500,
                    key=edited_key,
                )
                if edited_completion != asst_content:
                    st.caption("✏️ Edited — will save edited version on approve")

            with st.expander("Metadata"):
                st.json(meta)

            notes_key = f"notes_{i}_{pair_hash[:8]}"
            if notes_key not in st.session_state:
                st.session_state[notes_key] = ""

            notes = st.text_input(
                "Notes (optional)",
                key=notes_key,
                placeholder="Reviewer notes...",
            )

            # E3 — four-tier approval buttons
            _hint_key = f"hint_{i}_{pair_hash[:8]}"
            if _hint_key not in st.session_state:
                st.session_state[_hint_key] = ""

            _bc1, _bc2, _bc3, _bc4 = st.columns([1, 1.4, 1.4, 1])
            _disabled = not st.session_state.reviewer_id

            with _bc1:
                if st.button(
                    "✅ Approve",
                    key=f"approve_{i}_{pair_hash[:8]}",
                    disabled=_disabled,
                    help="Full approve — goes straight to JSONL",
                ):
                    # Save edited completion back into pair before approving
                    _edited = st.session_state.get(f"edit_{i}_{pair_hash[:8]}", "")
                    _entry_msgs = entry.get("pair", {}).get("messages", [])
                    for _m in _entry_msgs:
                        if _m["role"] == "assistant" and _edited:
                            _m["content"] = _edited
                    if approve(pair_hash, st.session_state.reviewer_id, notes, _QUEUE_PATH):
                        st.success("Approved!")
                        st.session_state.queue_refresh += 1
                        st.rerun()
                    else:
                        st.error("Approve failed")

            with _bc2:
                if st.button(
                    "✅ Conditional",
                    key=f"cond_{i}_{pair_hash[:8]}",
                    disabled=(_disabled or not notes),
                    help="Approve with caveat — note required",
                ):
                    if conditional_approve(pair_hash, st.session_state.reviewer_id,
                                           notes, _QUEUE_PATH):
                        st.success("Conditionally approved")
                        st.session_state.queue_refresh += 1
                        st.rerun()
                    else:
                        st.error("Conditional approve failed")
                if not notes and not _disabled:
                    st.caption("Add a note to use conditional approve")

            with _bc3:
                st.text_input(
                    "Correction hint",
                    key=_hint_key,
                    placeholder="e.g. Finding 1 should cite SOP §3.2",
                )
                if st.button(
                    "🔄 Send for correction",
                    key=f"correct_{i}_{pair_hash[:8]}",
                    disabled=(_disabled or not st.session_state[_hint_key]),
                    help="Gemma re-runs with your hint",
                ):
                    if send_for_correction(pair_hash, st.session_state.reviewer_id,
                                           st.session_state[_hint_key], _QUEUE_PATH):
                        st.info("Sent for correction")
                        st.session_state.queue_refresh += 1
                        st.rerun()
                    else:
                        st.error("Send for correction failed")

            with _bc4:
                if st.button(
                    "❌ Reject",
                    key=f"reject_{i}_{pair_hash[:8]}",
                    disabled=_disabled,
                    help="Discard — not written to JSONL",
                ):
                    if reject(pair_hash, st.session_state.reviewer_id, notes, _QUEUE_PATH):
                        st.warning("Rejected.")
                        st.session_state.queue_refresh += 1
                        st.rerun()
                    else:
                        st.error("Reject failed")

            if not reviewer_id:
                st.caption("⚠️ Enter your reviewer ID above to use review actions")

st.divider()

# ===========================================================================
# SECTION 3 — Export
# ===========================================================================

st.header("3 · Export")

export_col1, export_col2 = st.columns([1, 2])

with export_col1:
    if st.button("📤 Write approved pairs", type="primary"):
        with st.spinner("Writing approved pairs to JSONL..."):
            export_result = write_approved(
                queue_path=_QUEUE_PATH,
                data_dir=_DATA_DIR,
            )
        st.session_state.last_export = export_result

if st.session_state.last_export is not None:
    exp = st.session_state.last_export
    st.success(
        f"Written: **{exp['written']}** · "
        f"Skipped (duplicates): {exp['skipped']} · "
        f"Errors: {exp['errors']}"
    )

st.subheader("Output files")

for stage, label in [("stage2", "Stage 2 — Domain"), ("stage3", "Stage 3 — Firm")]:
    try:
        fpath = get_output_path(stage, _DATA_DIR)
        n = count(fpath)
        dcol1, dcol2 = st.columns([3, 1])
        dcol1.markdown(f"**{label}** — `{fpath.name}` — {n} pairs")

        if fpath.exists() and n > 0:
            with open(fpath, "rb") as f:
                dcol2.download_button(
                    label="⬇️ Download",
                    data=f.read(),
                    file_name=fpath.name,
                    mime="application/jsonl",
                    key=f"dl_{stage}",
                )
        else:
            dcol2.caption("No data yet")
    except Exception:
        pass

# ===========================================================================
# SECTION 4 — NPO-CX-1.1 Engagement Workpaper Generator
# ===========================================================================

st.header("4 · Engagement Workpaper Generator (NPO-CX-1.1)")

st.caption(
    "Generate a filled NPO-CX-1.1 Engagement Acceptance & Continuance form "
    "from a client's prior-year audit report PDF using the SOP-driven rule "
    "engine and PDF section detector."
)

# --- Source mode (existing engagement vs. new client upload) ---
wp_mode = st.radio(
    "Source",
    options=["Existing engagement", "New client (upload PDF)"],
    horizontal=True,
    key="wp_mode",
)

engagement_options: list[str] = []
if _ENGAGEMENTS_DIR.is_dir():
    # Only require a PDF — the renderer always uses the canonical blank,
    # so any .docx in the folder (historical workpaper, partially filled form, etc.)
    # is intentionally ignored.
    engagement_options = sorted(
        p.name for p in _ENGAGEMENTS_DIR.iterdir()
        if p.is_dir() and list(p.glob("*.pdf"))
    )

engagement_dir: Path | None = None
new_client_name: str = ""
new_client_pdf_bytes: bytes | None = None
new_client_pdf_filename: str | None = None

col_sel1, col_sel2 = st.columns([2, 2])

with col_sel1:
    if wp_mode == "Existing engagement":
        if engagement_options:
            selected_engagement = st.selectbox(
                "Engagement",
                options=engagement_options,
                help="Subfolder under 'Engagement Accept and Cont Form/' containing the client PDF + workpaper template.",
                key="wp_engagement",
            )
            engagement_dir = _ENGAGEMENTS_DIR / selected_engagement
        else:
            st.info(
                "No existing engagements found. Switch to 'New client' to upload a PDF."
            )
    else:
        new_client_name = st.text_input(
            "Client name",
            help="A folder will be created under 'Engagement Accept and Cont Form/<ClientName>/'.",
            key="wp_new_client_name",
        )
        uploaded_pdf = st.file_uploader(
            "Upload PY audit report PDF",
            type=["pdf"],
            key="wp_new_pdf",
        )
        if uploaded_pdf is not None:
            new_client_pdf_bytes = uploaded_pdf.getvalue()
            new_client_pdf_filename = uploaded_pdf.name

with col_sel2:
    engagement_type_label = st.radio(
        "Engagement type",
        options=["Recurring", "Initial / 1st year"],
        horizontal=True,
        key="wp_engagement_type",
        help="Recurring engagements skip Part II (auto-N/A). Initial fills Part II with SOP defaults.",
    )

engagement_type = (
    ENGAGEMENT_RECURRING if engagement_type_label == "Recurring" else ENGAGEMENT_INITIAL
)

if engagement_dir is not None:
    pdfs = list(engagement_dir.glob("*.pdf"))
    if pdfs:
        st.caption(
            f"PDF: `{pdfs[0].name}` · "
            f"Template: canonical blank ({CANONICAL_BLANK_TEMPLATE.name})"
        )
elif wp_mode == "New client (upload PDF)" and new_client_pdf_filename:
    st.caption(
        f"PDF: `{new_client_pdf_filename}` · "
        f"Template: canonical blank ({CANONICAL_BLANK_TEMPLATE.name})"
    )

# --- Auditor inputs (optional overrides for needs_input fields) ---
if True:
    with st.expander("Auditor inputs (optional — leave blank to keep as needs_input)"):
        wp_col1, wp_col2 = st.columns(2)
        with wp_col1:
            in_completed_by = st.text_input("Completed by", key="wp_completed_by")
            in_completion_date = st.text_input("Completion date (e.g., 03/09/2026)", key="wp_completion_date")
            in_fye_date = st.text_input(
                "Statement of Financial Position Date (current engagement, e.g., 06/30/2025)",
                key="wp_fye_date",
            )
            in_signoff_date = st.text_input("Engagement partner sign-off date", key="wp_signoff_date")
            in_concurring = st.text_input("Concurring partner (optional)", key="wp_concurring")
        with wp_col2:
            in_q1a = st.selectbox(
                "Q1(a) Basis of accounting",
                options=["", "Accrual / GAAP", "Cash Basis", "Other"],
                key="wp_q1a",
            )
            in_q1b = st.selectbox(
                "Q1(b) Grant compliance",
                options=["", "2 CFR Part 200 – Uniform Guidance", "Other"],
                key="wp_q1b",
            )
            in_q1f = st.selectbox(
                "Q1(f) Federal tax / info return",
                options=["", "No", "Yes — Form 990", "Yes — Form 990-PF", "Yes — Form 199", "Yes — Form RRF-1"],
                key="wp_q1f",
            )
            in_q2j_remark = st.text_area(
                "Q2(j) Predecessor auditor remark (Initial engagements only)",
                key="wp_q2j_remark",
                height=80,
            )
            in_org_override = st.text_input(
                "Organization name (override detector hint)",
                key="wp_org_override",
            )

    # Decide whether the user has provided enough input to generate.
    is_existing_ready = (wp_mode == "Existing engagement" and engagement_dir is not None)
    is_new_ready = (
        wp_mode == "New client (upload PDF)"
        and bool(new_client_name.strip())
        and new_client_pdf_bytes is not None
    )
    generate_enabled = is_existing_ready or is_new_ready

    if st.button(
        "🧾 Generate workpaper",
        type="primary",
        key="wp_generate",
        disabled=not generate_enabled,
        help=(
            None if generate_enabled
            else "Provide client name + PDF (new client) or pick an existing engagement."
        ),
    ):
        auditor_inputs: dict = {}
        if in_completed_by: auditor_inputs["completed_by"] = in_completed_by
        if in_completion_date: auditor_inputs["completion_date"] = in_completion_date
        if in_fye_date: auditor_inputs["financial_position_date"] = in_fye_date
        if in_signoff_date: auditor_inputs["sign_off_date"] = in_signoff_date
        if in_concurring: auditor_inputs["concurring_partner"] = in_concurring
        if in_q1a: auditor_inputs["q1_a"] = in_q1a
        if in_q1b: auditor_inputs["q1_b"] = in_q1b
        if in_q1f: auditor_inputs["q1_f"] = in_q1f
        if in_q2j_remark: auditor_inputs["pii_q2_j"] = in_q2j_remark
        if in_org_override: auditor_inputs["organization_name"] = in_org_override

        # If new client, register first (creates folder + seeds PDF + blank template)
        if wp_mode == "New client (upload PDF)":
            try:
                engagement_dir = register_new_client(
                    client_name=new_client_name.strip(),
                    pdf_bytes=new_client_pdf_bytes,
                    pdf_filename=new_client_pdf_filename,
                )
                st.info(f"Registered new engagement folder: `{engagement_dir}`")
            except (FileNotFoundError, ValueError) as e:
                st.error(f"Could not register client: {e}")
                st.stop()

        with st.spinner("Running detector + rule engine + renderer..."):
            result = run_engagement(
                engagement_dir=engagement_dir,
                engagement_type=engagement_type,
                auditor_inputs=auditor_inputs,
            )
            output_path = render_workpaper(result)

        st.session_state.workpaper_result = result
        st.session_state.workpaper_output_path = str(output_path)

    # --- Results display ---
    if st.session_state.workpaper_result is not None:
        result = st.session_state.workpaper_result
        output_path = Path(st.session_state.workpaper_output_path)

        st.success(f"Workpaper generated: `{output_path.name}`")

        # Summary metrics
        summary = result.summary
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Resolved", summary.get("resolved", 0))
        m2.metric("Needs input", summary.get("needs_input", 0))
        m3.metric("N/A (Part II)", summary.get("na", 0))
        m4.metric("Total fields", sum(summary.values()))

        # Section detection
        st.subheader("PDF section detection")
        det_cols = st.columns(3)
        for col, key, label in [
            (det_cols[0], "sefa", "SEFA (Q1c)"),
            (det_cols[1], "supplementary", "Supplementary Info (Q1d)"),
            (det_cols[2], "compliance", "Compliance Section (Q2)"),
        ]:
            sec = result.detection.sections.get(key)
            if sec and sec.found:
                col.success(f"✓ {label} — found in {sec.location}")
            else:
                col.info(f"✗ {label} — not present (rule fallback applies)")

        # Key resolved values
        st.subheader("Key resolved fields")
        q_rows = []
        for fid in ["q1_c", "q1_d", "q2", "acceptance_decision"]:
            rf = result.resolved.get(fid)
            if rf:
                q_rows.append({
                    "Field": fid,
                    "Value": rf.value,
                    "Source": rf.source,
                    "SOP citation": rf.citation.get("sop") or "",
                    "PDF citation": rf.citation.get("pdf") or "",
                })
        st.dataframe(q_rows, use_container_width=True, hide_index=True)

        # Download buttons
        st.subheader("Downloads")
        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            with open(output_path, "rb") as f:
                dl_col1.download_button(
                    label="⬇️ Filled workpaper (.docx)",
                    data=f.read(),
                    file_name=output_path.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="wp_dl_docx",
                )
        trace_path = Path(result.output_dir) / "trace.json"
        if trace_path.exists():
            with open(trace_path, "rb") as f:
                dl_col2.download_button(
                    label="⬇️ Trace JSON (audit trail)",
                    data=f.read(),
                    file_name=f"{result.engagement_name}_trace.json",
                    mime="application/json",
                    key="wp_dl_trace",
                )

        # Full field breakdown
        with st.expander("Full field-by-field breakdown"):
            from dataclasses import asdict
            rows = []
            for fid, rf in result.resolved.items():
                rows.append({
                    "Field": fid,
                    "Status": rf.status,
                    "Value": (rf.value or "")[:80],
                    "Source": rf.source,
                    "SOP": rf.citation.get("sop") or "",
                    "PDF": (rf.citation.get("pdf") or "")[:60],
                })
            st.dataframe(rows, use_container_width=True, hide_index=True)

# ===========================================================================
# SECTION 5 — RAG + Rule Engine Playground (inspect / validate the approach)
# ===========================================================================

st.header("5 · RAG + Rule Engine Playground")

st.caption(
    "Inspect how the PDF section detector and rule engine behave for any "
    "client PDF. Use the what-if toggles to test SOP Table 3 acceptance "
    "logic. This panel does not produce a .docx — use Section 4 for that."
)

_BLOCKER_LABELS = {
    "q4": "Q4 — Unacceptable financial reporting framework",
    "q5_a": "Q5(a) — Management refuses responsibility",
    "q5_b": "Q5(b) — Management refuses access",
    "q5_c": "Q5(c) — Management refuses representations",
    "q9": "Q9 — Firm lacks independence (GAAS)",
    "q9_a": "Q9(a) — Yellow Book independence impaired",
    "q9_b": "Q9(b) — Indirect cost plan + recoveries >$1M",
    "q11": "Q11 — Management integrity in doubt",
}

# --- Inputs ---
pg_col1, pg_col2 = st.columns([1, 1])

with pg_col1:
    pg_source = st.radio(
        "PDF source",
        options=["Existing engagement", "Upload PDF"],
        horizontal=True,
        key="pg_source",
    )

    pdf_for_test: Path | None = None
    pdf_label: str | None = None

    if pg_source == "Existing engagement":
        engagement_dirs = (
            sorted(p.name for p in _ENGAGEMENTS_DIR.iterdir() if p.is_dir() and list(p.glob("*.pdf")))
            if _ENGAGEMENTS_DIR.is_dir() else []
        )
        if engagement_dirs:
            pg_eng = st.selectbox("Engagement", engagement_dirs, key="pg_engagement")
            pdfs = list((_ENGAGEMENTS_DIR / pg_eng).glob("*.pdf"))
            pdf_for_test = pdfs[0] if pdfs else None
            pdf_label = f"{pg_eng}/{pdfs[0].name}" if pdfs else None
        else:
            st.info("No engagement subfolders with PDFs found.")
    else:
        uploaded = st.file_uploader("Upload PY audit report PDF", type=["pdf"], key="pg_upload")
        if uploaded is not None:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(uploaded.read())
            tmp.close()
            pdf_for_test = Path(tmp.name)
            pdf_label = uploaded.name

with pg_col2:
    pg_engagement_type_label = st.radio(
        "Engagement type",
        options=["Recurring", "Initial / 1st year"],
        horizontal=True,
        key="pg_engagement_type",
    )
    pg_engagement_type = (
        ENGAGEMENT_RECURRING if pg_engagement_type_label == "Recurring" else ENGAGEMENT_INITIAL
    )

st.markdown("**What-if blocker toggles** — flip any to Yes to see acceptance flip to DO NOT ACCEPT")
toggle_cols = st.columns(4)
blocker_overrides: dict[str, str] = {}
for i, (fid, label) in enumerate(_BLOCKER_LABELS.items()):
    col = toggle_cols[i % 4]
    if col.checkbox(label, key=f"pg_blocker_{fid}"):
        blocker_overrides[fid] = "Yes"

if st.button("🔍 Run detector + rule engine", type="primary", key="pg_run", disabled=pdf_for_test is None):
    with st.spinner("Running..."):
        detection = detect_sections(pdf_for_test)
        # The rule engine needs an org name hint to mirror Section 4 behavior
        auditor_inputs = {}
        if detection.header_hints.get("organization_name"):
            auditor_inputs["organization_name"] = detection.header_hints["organization_name"]

        resolved = resolve_workpaper(
            pdf_path=pdf_for_test,
            engagement_type=pg_engagement_type,
            auditor_inputs=auditor_inputs,
        )

        # Apply what-if overrides AFTER baseline resolution
        for fid, override_value in blocker_overrides.items():
            if fid in resolved:
                resolved[fid].value = override_value
                resolved[fid].rule_applied = (resolved[fid].rule_applied or "") + " [overridden via playground]"

        # Re-compute acceptance_decision with the overrides applied
        _BLOCKER_FIDS = list(_BLOCKER_LABELS.keys())
        tripped = next((fid for fid in _BLOCKER_FIDS if resolved.get(fid) and resolved[fid].value == "Yes"), None)
        if "acceptance_decision" in resolved:
            if tripped:
                resolved["acceptance_decision"].value = "DO NOT ACCEPT / CONTINUE"
                resolved["acceptance_decision"].rule_applied = f"Blocked: {tripped} == Yes"
            else:
                resolved["acceptance_decision"].value = "ACCEPT / CONTINUE"
                resolved["acceptance_decision"].rule_applied = "All blocker fields evaluate to No"

    st.session_state.playground_detection = detection
    st.session_state.playground_resolved = resolved
    st.session_state.playground_pdf_label = pdf_label

# --- Outputs ---
if st.session_state.playground_resolved is not None:
    detection = st.session_state.playground_detection
    resolved = st.session_state.playground_resolved
    pdf_label = st.session_state.playground_pdf_label

    st.caption(f"Last run against: `{pdf_label}` ({pg_engagement_type_label})")

    # A. Detector panel
    st.subheader("A · PDF Section Detector")

    det_cols = st.columns(3)
    for col, key, label in [
        (det_cols[0], "sefa", "SEFA → Q1(c)"),
        (det_cols[1], "supplementary", "Supplementary Info → Q1(d)"),
        (det_cols[2], "compliance", "Compliance Section → Q2 ref"),
    ]:
        sec = detection.sections.get(key)
        if sec and sec.found:
            col.success(f"✓ {label}")
            col.caption(f"matched: `{sec.match_text}` ({sec.location})")
        else:
            col.error(f"✗ {label}")
            col.caption("section not present — rule fallback applies")

    hint_cols = st.columns(3)
    hint_cols[0].metric("Pages", detection.page_count)
    hint_cols[1].metric("Extraction quality", detection.extraction_quality)
    hint_cols[2].metric("Chars extracted", f"{detection.total_chars_extracted:,}")

    h_col1, h_col2 = st.columns(2)
    h_col1.text(f"Org name hint: {detection.header_hints.get('organization_name') or '(none)'}")
    h_col2.text(f"PY FYE hint: {detection.header_hints.get('prior_year_fye_date') or '(none)'}")

    # B. Rule engine panel
    st.subheader("B · Rule Engine Output")

    from workpaper_generator.rule_engine import summarize as _summarize
    s = _summarize(resolved)
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total", sum(s.values()))
    m2.metric("Resolved", s.get("resolved", 0))
    m3.metric("Needs input", s.get("needs_input", 0))
    m4.metric("N/A", s.get("na", 0))

    decision_rf = resolved.get("acceptance_decision")
    if decision_rf:
        if decision_rf.value == "ACCEPT / CONTINUE":
            st.success(f"✅  **{decision_rf.value}** — {decision_rf.rule_applied}")
        else:
            st.error(f"⛔  **{decision_rf.value}** — {decision_rf.rule_applied}")

    pin_cols = st.columns(3)
    for col, fid, label in [
        (pin_cols[0], "q1_c", "Q1(c) Single Audit"),
        (pin_cols[1], "q1_d", "Q1(d) Supplementary Info"),
        (pin_cols[2], "q2", "Q2 Non-attest services"),
    ]:
        rf = resolved.get(fid)
        if rf:
            col.metric(label, rf.value or "—")
            col.caption(rf.citation.get("pdf") or rf.citation.get("sop") or "")

    # C. Full field table (filterable)
    st.subheader("C · All Resolved Fields")

    f_col1, f_col2 = st.columns([1, 1])
    source_filter = f_col1.multiselect(
        "Filter by source",
        options=sorted({rf.source for rf in resolved.values()}),
        default=[],
        key="pg_source_filter",
    )
    status_filter = f_col2.multiselect(
        "Filter by status",
        options=sorted({rf.status for rf in resolved.values()}),
        default=[],
        key="pg_status_filter",
    )

    table_rows = []
    for fid, rf in resolved.items():
        if source_filter and rf.source not in source_filter:
            continue
        if status_filter and rf.status not in status_filter:
            continue
        table_rows.append({
            "Field": fid,
            "Status": rf.status,
            "Value": (rf.value or "")[:80],
            "Source": rf.source,
            "Rule applied": (rf.rule_applied or "")[:60],
            "SOP citation": rf.citation.get("sop") or "",
            "PDF citation": (rf.citation.get("pdf") or "")[:60],
        })
    st.dataframe(table_rows, use_container_width=True, hide_index=True)
    st.caption(f"Showing {len(table_rows)} of {len(resolved)} fields")