"""
auditai_data_normalization/extractors/docx_extractor.py
=========================================================
Extracts text, tables, and structure from .docx files using python-docx.

Why python-docx and not an LLM or OCR?
---------------------------------------
DOCX is a ZIP of XML. python-docx reads it deterministically — zero
hallucination risk, preserves paragraph styles (Heading1/2/3, Normal,
Table), and is the fastest extractor in the pipeline. Every workpaper
that arrives as .docx (PPC forms, engagement letters, management letters)
goes through this extractor first.

What it produces
-----------------
- sections  : one Section per heading block (heading + following paragraphs)
              If no headings are found, one Section per non-empty paragraph.
- tables    : one ExtractedTable per Word table, with headers and rows.
              Merged cells are forward-filled so every row has same length.
- metadata  : heading_count, table_count, paragraph_count,
              has_tracked_changes, style_names_found

Edge cases handled
------------------
- Merged cells in tables (forward-fill)
- Tracked changes (LibreOffice accept-all is recommended upstream;
  here we read both original and revision text and note the flag)
- Password-protected files (raises PasswordProtectedError)
- Legacy .doc files (raises UnsupportedFormatError — convert to .docx first)
- Empty documents (returns record with extraction_status='partial')
- Tables with no header row (headers set to col_0, col_1, ...)

Public API
----------
    extract(file_path: str | Path) -> DocumentRecord
"""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from auditai_data_normalization.schema import (
    DocumentRecord,
    ExtractedTable,
    Section,
)
from auditai_data_normalization.text_normalizer import normalize_text
from auditai_data_normalization.doc_classifier import detect_category

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------

class PasswordProtectedError(Exception):
    """Raised when the .docx file is encrypted and cannot be opened."""

class UnsupportedFormatError(Exception):
    """Raised for .doc (legacy binary) files — convert to .docx first."""


# ---------------------------------------------------------------------------
# Heading style detection
# ---------------------------------------------------------------------------

# python-docx style names for headings (covers PPC templates)
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3", "heading 4",
    "heading 5", "heading 6",
    # PPC-specific style variants
    "h1", "h2", "h3",
    "title", "subtitle",
}

def _is_heading(para: Paragraph) -> bool:
    """True if paragraph has a heading style."""
    if para.style and para.style.name:
        return para.style.name.lower() in _HEADING_STYLES
    return False


# ---------------------------------------------------------------------------
# Table extraction helpers
# ---------------------------------------------------------------------------

def _get_cell_text(cell) -> str:
    """
    Extract full text from a table cell, joining paragraphs with spaces.
    Phase 1: normalize_text() applied before returning so checkbox symbols,
    underscore artifacts, and inline markers are resolved at the cell level.
    """
    raw = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
    return normalize_text(raw)
 

# ---------------------------------------------------------------------------
# Structural parsing helpers (Phase 3)
# ---------------------------------------------------------------------------
 
# Known header labels that should NOT be promoted as values
# during vertical continuation. These appear as row-0 labels in
# PPC forms and would be misread as values for the row above.
_KNOWN_HEADER_LABELS: frozenset[str] = frozenset({
    "yes", "no", "n/a", "not applicable", "check all that apply",
    "amount", "description", "explanation", "comments", "instructions",
    "col_0", "col_1", "col_2", "col_3", "col_4",
})
 
# Checkbox sentinels produced by Phase 1 normalize_text()
_CHECKED_SENTINEL   = "true"
_UNCHECKED_SENTINEL = "false"
 
 
def _looks_like_value(cell_text: str) -> bool:
    """
    Heuristic: does this cell look like a data value rather than a label?
 
    Used by vertical continuation to decide whether the next row's label
    cell should be promoted as the value for the current row.
 
    A cell looks like a value when:
    - It has content AND
    - It does NOT end with ":" AND
    - It is NOT a known header label AND
    - It is NOT a pure checkbox sentinel ("true" / "false")
      unless the column is a checkbox column (handled separately)
    """
    t = cell_text.strip()
    if not t:
        return False
    if t.lower() in _KNOWN_HEADER_LABELS:
        return False
    if t.endswith(":"):
        return False
    # Pure checkbox sentinels are values only in the checkbox column pass
    if t.lower() in (_CHECKED_SENTINEL, _UNCHECKED_SENTINEL):
        return False
    return True
 
 
def _resolve_checkbox_columns(
    headers: list[str],
    data_rows: list[list[str]],
) -> list[tuple[str, str]]:
    """
    Scan table for checkbox columns and return (field_label, value) pairs.
 
    Phase 3 — checkbox column tracking.
 
    Strategy:
    1. Find columns whose header is a known Yes/No/checkbox label.
    2. For each data row, find which of those columns contains a checked
       sentinel ("true" from Phase 1 normalization).
    3. The row's label (col 0) + the matched column header → field: value.
 
    This resolves tables like:
 
        Label              | Yes | No
        Single Audit       | ✔   | ☐       →  includes_single_audit: true
        GAGAS Audit        | ☐   | ✔       →  includes_gagas: false
 
    Parameters
    ----------
    headers : list[str]
        Column headers (already normalized by Phase 1).
    data_rows : list[list[str]]
        Table data rows (already normalized by Phase 1).
 
    Returns
    -------
    list[tuple[str, str]]
        (label, resolved_value) pairs ready to be written to raw_text
        as \"label: value\" lines for normalize.py field extraction.
    """
    if not headers or not data_rows:
        return []
 
    # Identify Yes/No/checkbox column indices
    _YES_VARIANTS  = {"yes", "true",  "✓", "applicable", "check", "checked"}
    _NO_VARIANTS   = {"no",  "false", "✗", "not applicable", "n/a", "unchecked"}
 
    yes_col_indices: list[int] = []
    no_col_indices:  list[int] = []
 
    for col_idx, header in enumerate(headers):
        h = header.strip().lower()
        if h in _YES_VARIANTS:
            yes_col_indices.append(col_idx)
        elif h in _NO_VARIANTS:
            no_col_indices.append(col_idx)
 
    if not yes_col_indices and not no_col_indices:
        return []   # No checkbox columns found — nothing to do
 
    pairs: list[tuple[str, str]] = []
 
    for row in data_rows:
        if not row:
            continue
        # Label is always the first non-empty cell in the row
        row_label = next((c.strip() for c in row if c.strip()), "")
        if not row_label or row_label.lower() in _KNOWN_HEADER_LABELS:
            continue
 
        resolved_value: str | None = None
 
        # Check Yes columns first
        for col_idx in yes_col_indices:
            if col_idx < len(row):
                cell = row[col_idx].strip().lower()
                if cell == _CHECKED_SENTINEL:
                    resolved_value = "true"
                    break
                elif cell == _UNCHECKED_SENTINEL:
                    resolved_value = "false"
                    break
 
        # If not resolved from Yes col, check No columns
        if resolved_value is None:
            for col_idx in no_col_indices:
                if col_idx < len(row):
                    cell = row[col_idx].strip().lower()
                    if cell == _CHECKED_SENTINEL:
                        resolved_value = "false"   # "No" column is checked → value is false
                        break
                    elif cell == _UNCHECKED_SENTINEL:
                        resolved_value = "true"    # "No" column is unchecked → value is true
                        break
 
        if resolved_value is not None:
            pairs.append((row_label, resolved_value))
 
    return pairs


def _extract_table(table: Table, index: int, source: str = "") -> ExtractedTable:
    """
    Convert a python-docx Table object to an ExtractedTable.
 
    Phase 3 additions on top of the existing forward-fill logic:
    A. Vertical continuation — if a row has a label but empty value,
       and the next row's label cell looks like a value, promote it.
    B. Checkbox column tracking — detect Yes/No columns and emit
       structured label: true/false lines into raw_text.
    """
    raw_rows: list[list[str]] = []
    for row in table.rows:
        cells = [_get_cell_text(cell) for cell in row.cells]
        raw_rows.append(cells)
 
    if not raw_rows:
        return ExtractedTable(index=index, source=source)
 
    # ── Horizontal forward-fill (existing logic, unchanged) ──────────────
    filled_rows: list[list[str]] = []
    for row in raw_rows:
        filled: list[str] = []
        prev = ""
        for i, cell in enumerate(row):
            if cell == prev and i > 0:
                filled.append("")
            else:
                filled.append(cell)
                prev = cell
        filled_rows.append(filled)
 
    # ── Phase 3A — Vertical continuation ────────────────────────────────
    # For two-column label/value tables: if row[i] has a label (col 0)
    # but an empty value (col 1), check if row[i+1]'s col 0 looks like
    # a value — if so, promote it and mark row[i+1] as consumed.
    #
    # Only applied to tables with at least 2 columns.
    # Guards against consuming genuine label rows.
 
    if len(filled_rows) > 1 and filled_rows and len(filled_rows[0]) >= 2:
        consumed: set[int] = set()
        continued_rows: list[list[str]] = []
 
        for i, row in enumerate(filled_rows):
            if i in consumed:
                continue
 
            # Two-column check: label in col 0, value in col 1
            label_cell = row[0].strip() if len(row) > 0 else ""
            value_cell = row[1].strip() if len(row) > 1 else ""
 
            if label_cell and not value_cell and i + 1 < len(filled_rows):
                next_row  = filled_rows[i + 1]
                next_label = next_row[0].strip() if next_row else ""
                next_value = next_row[1].strip() if len(next_row) > 1 else ""
 
                # Promote if next row's col 0 looks like a value
                # AND next row has no value in col 1 of its own
                if _looks_like_value(next_label) and not next_value:
                    # Merge: current row label + next row label as value
                    merged = list(row)
                    merged[1] = next_label
                    continued_rows.append(merged)
                    consumed.add(i + 1)
                    continue
 
            continued_rows.append(row)
 
        filled_rows = continued_rows
 
    # ── Header detection (existing logic, unchanged) ─────────────────────
    headers: list[str] = []
    data_rows: list[list[str]] = filled_rows
 
    if len(filled_rows) >= 1:
        first_row = filled_rows[0]
        non_empty = [c for c in first_row if c.strip()]
        if non_empty:
            headers = first_row
            data_rows = filled_rows[1:] if len(filled_rows) > 1 else []
        else:
            headers = [f"col_{i}" for i in range(len(first_row))]
 
    # ── Phase 3B — Checkbox column tracking ─────────────────────────────
    # Detect Yes/No columns, resolve checkbox state per row, emit
    # structured label: value lines that field extraction can parse.
    checkbox_pairs = _resolve_checkbox_columns(headers, data_rows)
 
    # ── Build raw_text (existing logic + checkbox pairs appended) ────────
    lines: list[str] = []
    if headers:
        lines.append(" | ".join(h for h in headers if h))
        lines.append("-" * 40)
    for row in data_rows:
        if any(c.strip() for c in row):
            if headers:
                pairs = [
                    f"{h}: {v}"
                    for h, v in zip(headers, row)
                    if h.strip() and v.strip()
                ]
                lines.append(", ".join(pairs) if pairs else " | ".join(row))
            else:
                lines.append(" | ".join(c for c in row if c))
 
    # Append checkbox-resolved pairs as parseable label: value lines
    # These are picked up by normalize.py's _extract_fields_from_record()
    if checkbox_pairs:
        lines.append("# checkbox-resolved")
        for label, value in checkbox_pairs:
            lines.append(f"{label}: {value}")
 
    return ExtractedTable(
        index=index,
        source=source,
        headers=headers,
        rows=data_rows,
        raw_text="\\n".join(lines),
    )

   

# ---------------------------------------------------------------------------
# Tracked changes detection
# ---------------------------------------------------------------------------

def _has_tracked_changes(doc: Document) -> bool:
    """Return True if document contains any tracked change markup."""
    body = doc.element.body
    # Tracked changes use w:ins and w:del XML elements
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return (
        body.find(f"{{{ns}}}ins") is not None
        or body.find(f"{{{ns}}}del") is not None
    )


# ---------------------------------------------------------------------------
# Section builder
# ---------------------------------------------------------------------------

def _build_sections(doc: Document) -> list[Section]:
    """
    Walk document body elements and group them into Sections.

    Strategy:
    - Each heading paragraph starts a new Section.
    - All following non-heading paragraphs are collected into that Section.
    - Tables are rendered as text and appended to the current Section,
      AND returned separately in the tables list.
    - If no headings found, each non-empty paragraph becomes its own Section.
    """
    sections: list[Section] = []
    current_heading = ""
    current_lines: list[str] = []
    section_index = 0

    def _flush(heading: str, lines: list[str]) -> None:
        nonlocal section_index
        content = "\n".join(l for l in lines if l.strip())
        if content.strip() or heading.strip():
            sections.append(Section(
                index=section_index,
                heading=heading,
                content=content,
                token_count=len(content.split()),
            ))
            section_index += 1

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Paragraph
            para_text = "".join(
                node.text or ""
                for node in block.iter()
                if node.tag.endswith("}t")
            ).strip()

            # Check style
            style_name = ""
            pPr = block.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_name = (pStyle.get(qn("w:val")) or "").lower()

            is_hdg = style_name in _HEADING_STYLES or any(
                style_name.startswith(h) for h in ["heading", "h1", "h2", "h3"]
            )

            if is_hdg and para_text:
                _flush(current_heading, current_lines)
                current_heading = para_text
                current_lines = []
            elif para_text:
                current_lines.append(para_text)

        elif tag == "tbl":
            # Table — render as text block and add to current section
            # (Full ExtractedTable objects are built separately in extract())
            rows_text: list[str] = []
            for tr in block.findall(f".//{qn('w:tr')}"):
                cells_text = []
                for tc in tr.findall(f".//{qn('w:tc')}"):
                    cell_text = "".join(
                        node.text or ""
                        for node in tc.iter()
                        if node.tag.endswith("}t")
                    ).strip()
                    if cell_text:
                        cells_text.append(cell_text)
                if cells_text:
                    rows_text.append(" | ".join(cells_text))
            if rows_text:
                current_lines.append("\n".join(rows_text))

    # Flush final section
    _flush(current_heading, current_lines)

    # Fallback: if no sections produced, one section per paragraph
    if not sections:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                sections.append(Section(
                    index=i,
                    heading="",
                    content=para.text.strip(),
                    token_count=len(para.text.split()),
                ))

    return sections


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def extract(file_path: str | Path) -> DocumentRecord:
    """
    Extract content from a .docx file.

    Parameters
    ----------
    file_path : str | Path
        Path to a .docx file. Must exist and be readable.

    Returns
    -------
    DocumentRecord
        Populated with sections, tables, metadata, and provenance fields.
        pii_scrubbed=False — call pii.scrub_record() after this.
        auditor_approved=False — set after human review.

    Raises
    ------
    FileNotFoundError
        If file_path does not exist.
    UnsupportedFormatError
        If file is a legacy .doc binary (not .docx ZIP format).
    PasswordProtectedError
        If file is encrypted.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # Reject legacy .doc (binary OLE format)
    if path.suffix.lower() == ".doc":
        raise UnsupportedFormatError(
            f"{path.name} is a legacy .doc file. "
            "Convert to .docx via LibreOffice headless first: "
            "libreoffice --headless --convert-to docx <file>"
        )

    # Compute file hash before opening
    raw_bytes = path.read_bytes()
    file_hash = hashlib.sha256(raw_bytes).hexdigest()
    file_size = len(raw_bytes)

    # Detect password protection (encrypted DOCX starts with OLE magic bytes)
    # D0 CF 11 E0 = OLE2 compound document (encrypted Office file)
    if raw_bytes[:4] == b"\xd0\xcf\x11\xe0":
        raise PasswordProtectedError(
            f"{path.name} appears to be password-protected (OLE encryption). "
            "Supply the password via the intake API before extraction."
        )

    # Open document
    try:
        doc = Document(str(path))
    except Exception as e:
        return DocumentRecord(
            source_path=str(path),
            file_name=path.name,
            file_type="docx",
            file_size_bytes=file_size,
            file_hash=file_hash,
            extraction_method="python_docx",
            extraction_status="failed",
            extraction_error=str(e),
            needs_review=True,
        )

    # ---- Extract tables ------------------------------------------------
    tables: list[ExtractedTable] = []
    for i, tbl in enumerate(doc.tables):
        tables.append(_extract_table(tbl, index=i, source=f"Table {i + 1}"))

    # ---- Extract sections ----------------------------------------------
    sections = _build_sections(doc)

    # ---- Collect style names used (metadata) ---------------------------
    style_names = list({
        p.style.name
        for p in doc.paragraphs
        if p.style and p.style.name and p.text.strip()
    })

    # ---- Build cleaned_text from sections ------------------------------
    cleaned_text = "\n\n".join(
        (f"{s.heading}\n{s.content}" if s.heading else s.content).strip()
        for s in sections
        if s.content.strip() or s.heading.strip()
    )

    # ---- Raw text (unjoined, for dedup hashing) ------------------------
    raw_text = "\n".join(p.text for p in doc.paragraphs)

    # ---- Word count ----------------------------------------------------
    word_count = len(cleaned_text.split())

    # ---- Determine extraction status -----------------------------------
    if not sections and not tables:
        status = "failed"
        error = "No content extracted — document may be empty"
    elif not sections:
        status = "partial"
        error = "No text sections found — tables only"
    else:
        status = "success"
        error = ""

    # ---- Detect document category (Phase 2) ----------------------------
    doc_category = detect_category(
        file_name=path.name,
        sections=sections,
        cleaned_text=cleaned_text,
    )

    # ---- Metadata ------------------------------------------------------
    metadata = {
        "heading_count": sum(1 for s in sections if s.heading),
        "table_count": len(tables),
        "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
        "has_tracked_changes": _has_tracked_changes(doc),
        "style_names_found": style_names,
        "section_count": len(sections),
        "document_category": doc_category,       # Phase 2 — consumed by llm_extractor fallback
    }
 
    logger.info(
        "docx_extractor: %s — %d sections, %d tables, %d words, status=%s",
        path.name, len(sections), len(tables), word_count, status,
    )

    return DocumentRecord(
        source_path=str(path),
        file_name=path.name,
        file_type="docx",
        file_size_bytes=file_size,
        file_hash=file_hash,
        raw_text=raw_text,
        cleaned_text=cleaned_text,
        sections=sections,
        tables=tables,
        extraction_method="python_docx",
        extraction_status=status,
        extraction_error=error,
        extraction_confidence=0.0,   # set by confidence.py after all extractors run
        word_count=word_count,
        page_count=0,                # DOCX has no page count in the DOM
        ocr_used=False,
        pii_scrubbed=False,
        needs_review=(status != "success"),
        metadata=metadata,
    )