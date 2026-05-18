"""
auditai_data_normalization/tests/test_phase1.py
================================================
Phase 1 exit criteria tests.

Three scenarios required before Phase 2 starts
(from AuditAI_Engineering_Benchmark.pptx):

    Scenario 1 — score_confidence() on known-good/bad inputs
    Scenario 2 — 5 sample workpapers end-to-end through normalize_document()
    Scenario 3 — PII stripping verified (EIN, SSN, entity names confirmed gone)

Run with:
    pytest auditai_data_normalization/tests/test_phase1.py -v

All tests are self-contained — they create their own sample files
where real files are not available, so the suite runs on any machine
with the venv activated.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import sys
import tempfile
import types
from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Stub spaCy before any pipeline import so tests run without en_core_web_lg
# on machines where the model is not installed (CI, fresh clones).
# On the GPU machine where en_core_web_lg IS installed, the stub is never
# used because the real module loads first.
# ---------------------------------------------------------------------------
if "spacy" not in sys.modules:
    fake_spacy = types.ModuleType("spacy")
    fake_spacy.load = lambda *a, **kw: (_ for _ in ()).throw(OSError("no model"))
    fake_spacy.util = types.ModuleType("spacy.util")
    fake_spacy.util.get_installed_models = lambda: []
    sys.modules["spacy"] = fake_spacy


# Force regex PII tier — tests never need NER, keeps them fast and portable
from auditai_data_normalization import pii as _pii_module
_pii_module._presidio_analyzer = False
_pii_module._presidio_available = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SAMPLES_DIR = Path(__file__).parent / "samples"
_SAMPLES_DIR.mkdir(exist_ok=True)


def _make_simple_docx(path: Path, text_lines: list[str]) -> None:
    """Create a minimal DOCX with the given text lines."""
    from docx import Document
    doc = Document()
    for line in text_lines:
        doc.add_paragraph(line)
    doc.save(str(path))


def _make_simple_xlsx(path: Path, sheets: dict[str, list[list]]) -> None:
    """Create a minimal XLSX with given sheets."""
    import openpyxl
    wb = openpyxl.Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        ws = wb.active if first else wb.create_sheet(sheet_name)
        if first:
            ws.title = sheet_name
            first = False
        for row in rows:
            ws.append(row)
    wb.save(str(path))


def _make_simple_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, indent=2))


def _make_simple_csv(path: Path, rows: list[list]) -> None:
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        for row in rows:
            writer.writerow(row)


# ---------------------------------------------------------------------------
# SCENARIO 1 — score_confidence() on known-good / bad inputs
# ---------------------------------------------------------------------------

class TestScoreConfidence:
    """Scenario 1: score_confidence() scoring rules."""

    def test_all_three_agree_returns_1_0(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("client_name", "ECE STEP", "ECE STEP", "ECE STEP") == 1.0

    def test_two_of_three_agree_returns_0_7(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("client_name", "ECE STEP", "ECE STEP", "Other Name") == 0.7

    def test_all_three_disagree_returns_0_2(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("client_name", "A", "B", "C") == 0.2

    def test_all_null_returns_0_0(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("has_single_audit", None, None, None) == 0.0

    def test_two_ran_both_agree_returns_0_9(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("fiscal_year_end", "2025-06-30", "2025-06-30", None) == 0.9

    def test_two_ran_both_disagree_returns_0_3(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("fiscal_year_end", "2025-06-30", "2024-06-30", None) == 0.3

    def test_one_ran_has_value_returns_0_6(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("engagement_partner", "Jane Smith", None, None) == 0.6

    def test_one_ran_empty_returns_0_0(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("engagement_partner", "", None, None) == 0.0

    def test_casing_normalised_before_comparison(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("client_type", "NPO", "npo", "  NPO  ") == 1.0

    def test_whitespace_normalised_before_comparison(self):
        from auditai_data_normalization.confidence import score_confidence
        assert score_confidence("client_name", "ECE  STEP", "ECE STEP", "ECE STEP") == 1.0

    def test_score_fields_batch(self):
        from auditai_data_normalization.confidence import score_fields
        scores = score_fields({
            "client_type":    ["NPO", "NPO", "NPO"],
            "fiscal_year_end":["2025-06-30", "2025-06-30", None],
            "missing_field":  [None, None, None],
        })
        assert scores["client_type"] == 1.0
        assert scores["fiscal_year_end"] == 0.9
        assert scores["missing_field"] == 0.0

    def test_score_record_weighted_aggregate(self):
        from auditai_data_normalization.confidence import score_record
        # high-importance field missing → drags aggregate below 0.7
        scores = {
            "client_type":      1.0,
            "has_single_audit": 0.0,   # high-importance, missing
            "fiscal_year_end":  0.9,
        }
        agg = score_record(scores)
        assert agg < 0.7, f"Expected < 0.7, got {agg}"

    def test_summarise_gate_pass(self):
        from auditai_data_normalization.confidence import summarise
        scores = {"client_type": 1.0, "fiscal_year_end": 0.9, "client_name": 0.7}
        s = summarise(scores)
        assert s.passes_gate is True
        assert s.aggregate_score >= 0.7

    def test_summarise_gate_fail(self):
        from auditai_data_normalization.confidence import summarise
        scores = {"client_type": 0.2, "has_single_audit": 0.0}
        s = summarise(scores)
        assert s.passes_gate is False
        assert "has_single_audit" in s.fields_missing


# ---------------------------------------------------------------------------
# SCENARIO 2 — 5 workpapers end-to-end through normalize_document()
# ---------------------------------------------------------------------------

class TestNormalizeEndToEnd:
    """Scenario 2: 5 document types through the full pipeline."""

    # ------ 2a: DOCX ------

    def test_docx_engagement_form(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document

        docx_path = tmp_path / "engagement_form.docx"
        _make_simple_docx(docx_path, [
            "Engagement Acceptance and Continuance Form",
            "Organization: ABC Nonprofit LLC",
            "Statement of Financial Position Date: 06/30/2025",
            "Completed by: MS1    Date: 08/01/2025",
            "Engagement Partner: Jane Smith",
            "Single Audit: Not Applicable",
            "Audit of financial statements in accordance with GAAS: Yes",
        ])

        rec = normalize_document(str(docx_path), run_parallel=False)

        assert rec.extraction_status == "success"
        assert rec.file_type == "docx"
        assert rec.pii_scrubbed is True
        assert len(rec.sections) > 0
        assert rec.word_count > 0
        assert rec.file_hash != ""
        assert len(rec.file_hash) == 64
        assert rec.extraction_method == "python_docx"
        assert rec.auditor_approved is False   # always False until human approves

    def test_docx_record_is_ready_after_approval(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document
        from auditai_data_normalization.review_queue import approve_record

        docx_path = tmp_path / "form.docx"
        _make_simple_docx(docx_path, ["Client: Test Org Inc.", "FYE: 2025-06-30"])

        rec = normalize_document(str(docx_path), run_parallel=False)
        # Manually push confidence above gate for this test
        rec.extraction_confidence = 0.75
        rec.needs_review = False

        rec = approve_record(rec, reviewer_id="SH")
        assert rec.auditor_approved is True
        assert rec.reviewer_id == "SH"
        assert rec.review_date != ""
        assert rec.is_ready_for_training() is True

    # ------ 2b: PDF (text-native) ------

    def test_pdf_text_native(self, tmp_path):
        """Create a simple text PDF and verify extraction."""
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            pdf_path = tmp_path / "simple.pdf"
            c = rl_canvas.Canvas(str(pdf_path))
            c.drawString(72, 750, "FINANCIAL STATEMENTS")
            c.drawString(72, 730, "Total assets: $1,000,000")
            c.drawString(72, 710, "Net assets: $500,000")
            c.save()
            _run_pdf_test(str(pdf_path))
        except ImportError:
            pytest.skip("reportlab not installed — skipping PDF creation test")

    # ------ 2c: XLSX ------

    def test_xlsx_trial_balance(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document

        xlsx_path = tmp_path / "trial_balance.xlsx"
        _make_simple_xlsx(xlsx_path, {
            "Trial Balance": [
                ["Account Code", "Account Name", "Debit", "Credit"],
                ["1010", "Cash", "500000", ""],
                ["2010", "Accounts Payable", "", "100000"],
                ["3010", "Net Assets", "", "400000"],
                ["", "TOTAL", "500000", "500000"],
            ],
            "Cover": [["Instructions only"]],
        })

        rec = normalize_document(str(xlsx_path), run_parallel=False)

        assert rec.extraction_status == "success"
        assert rec.file_type == "xlsx"
        assert rec.pii_scrubbed is True
        assert len(rec.sections) == 1       # Cover skipped
        assert len(rec.tables) == 1
        assert rec.tables[0].headers[0] == "Account Code"
        assert rec.metadata["skipped_sheets"] == ["Cover"]
        assert "Trial Balance" in rec.metadata["numeric_sheets"]

    # ------ 2d: CSV ------

    def test_csv_simple(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document

        csv_path = tmp_path / "data.csv"
        _make_simple_csv(csv_path, [
            ["account_code", "account_name", "balance"],
            ["1010", "Cash", "505900"],
            ["1020", "Accounts Receivable", "380277"],
        ])

        rec = normalize_document(str(csv_path), run_parallel=False)

        assert rec.extraction_status == "success"
        assert rec.file_type == "csv"
        assert rec.pii_scrubbed is True
        assert rec.extraction_method == "pandas_csv"
        assert len(rec.tables) == 1
        assert rec.tables[0].headers == ["account_code", "account_name", "balance"]

    # ------ 2e: JSON ------

    def test_json_structured(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document

        json_path = tmp_path / "metadata.json"
        _make_simple_json(json_path, {
            "client_name": "Early Childhood Education STEP",
            "fiscal_year_end": "2025-06-30",
            "audit_type": "GAAS",
            "has_single_audit": False,
        })

        rec = normalize_document(str(json_path), run_parallel=False)

        assert rec.extraction_status == "success"
        assert rec.file_type == "json"
        assert rec.pii_scrubbed is True
        assert rec.extraction_method == "stdlib_json"
        assert len(rec.sections) == 4   # one per top-level key
        headings = [s.heading for s in rec.sections]
        assert "client_name" in headings
        assert "fiscal_year_end" in headings

    # ------ General contract tests ------

    def test_file_hash_is_sha256(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document

        p = tmp_path / "hash_test.json"
        _make_simple_json(p, {"key": "value"})
        rec = normalize_document(str(p), run_parallel=False)

        expected = hashlib.sha256(p.read_bytes()).hexdigest()
        assert rec.file_hash == expected

    def test_file_not_found_raises(self):
        from auditai_data_normalization.normalize import normalize_document
        with pytest.raises(FileNotFoundError):
            normalize_document("/tmp/this_does_not_exist_ever.docx")

    def test_skip_extension_raises(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document
        p = tmp_path / "index.cdx"
        p.write_bytes(b"fake")
        with pytest.raises(ValueError, match="skipped"):
            normalize_document(str(p))

    def test_unsupported_extension_raises(self, tmp_path):
        from auditai_data_normalization.normalize import normalize_document
        p = tmp_path / "file.xyz"
        p.write_bytes(b"fake content")
        with pytest.raises(ValueError):
            normalize_document(str(p))


def _run_pdf_test(pdf_path: str) -> None:
    """Shared PDF test logic."""
    from auditai_data_normalization.normalize import normalize_document
    rec = normalize_document(pdf_path, run_parallel=False)
    assert rec.extraction_status in ("success", "partial")
    assert rec.file_type in ("pdf_text", "pdf_scanned")
    assert rec.pii_scrubbed is True
    assert rec.file_hash != ""


# ---------------------------------------------------------------------------
# SCENARIO 3 — PII stripping verified
# ---------------------------------------------------------------------------

class TestPIIScrubbing:
    """Scenario 3: EIN, SSN, entity names confirmed stripped."""

    def test_ein_replaced(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("The EIN is 12-3456789 for this organization.")
        assert "12-3456789" not in result.cleaned_text
        assert "[EIN]" in result.cleaned_text

    def test_ssn_replaced(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("SSN: 123-45-6789 on file.")
        assert "123-45-6789" not in result.cleaned_text
        assert "[SSN]" in result.cleaned_text

    def test_itin_replaced(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("ITIN 912-34-5678 was provided.")
        assert "912-34-5678" not in result.cleaned_text

    def test_client_entity_replaced(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("Audit of Early Childhood Education LLC for FY2024.")
        assert "Early Childhood Education LLC" not in result.cleaned_text
        assert "[CLIENT_ENTITY]" in result.cleaned_text

    def test_account_num_replaced_with_context(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("Account 123456789012 balance.")
        assert "123456789012" not in result.cleaned_text
        assert "[ACCOUNT_NUM]" in result.cleaned_text

    def test_account_num_not_replaced_without_context(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("Reference number 123456789012 in the report.")
        # Should NOT replace — no banking context keyword
        assert "123456789012" in result.cleaned_text

    def test_multiple_pii_types_in_one_document(self):
        from auditai_data_normalization.pii import scrub
        text = (
            "Client: ABC Nonprofit Inc. "
            "EIN: 94-1234567  SSN: 987-65-4321"
        )
        result = scrub(text)
        assert "94-1234567" not in result.cleaned_text
        assert "987-65-4321" not in result.cleaned_text
        types_found = result.types_found()
        assert "EIN" in types_found
        assert "SSN" in types_found

    def test_scrub_record_sets_pii_scrubbed_flag(self, tmp_path):
        from auditai_data_normalization.schema import DocumentRecord, Section
        from auditai_data_normalization.pii import scrub_record

        rec = DocumentRecord(
            file_name="test.docx",
            file_type="docx",
            cleaned_text="EIN: 12-3456789 for Client LLC",
            sections=[
                Section(index=0, heading="Header",
                        content="SSN: 123-45-6789 on file")
            ],
        )
        assert rec.pii_scrubbed is False
        rec = scrub_record(rec)
        assert rec.pii_scrubbed is True
        assert "12-3456789" not in rec.cleaned_text
        assert "123-45-6789" not in rec.sections[0].content

    def test_raw_text_never_scrubbed(self, tmp_path):
        """raw_text must be preserved as-is for dedup hashing."""
        from auditai_data_normalization.schema import DocumentRecord
        from auditai_data_normalization.pii import scrub_record

        original_raw = "EIN: 12-3456789 raw content"
        rec = DocumentRecord(
            file_name="test.docx",
            file_type="docx",
            raw_text=original_raw,
            cleaned_text="EIN: 12-3456789 clean content",
        )
        rec = scrub_record(rec)
        # raw_text must be unchanged
        assert rec.raw_text == original_raw

    def test_pii_redaction_log_populated(self):
        from auditai_data_normalization.schema import DocumentRecord, Section
        from auditai_data_normalization.pii import scrub_record

        rec = DocumentRecord(
            file_name="test.docx",
            file_type="docx",
            cleaned_text="EIN: 12-3456789",
            sections=[Section(index=0, content="SSN: 987-65-4321")],
        )
        rec = scrub_record(rec)
        types = [r.pii_type for r in rec.pii_redactions]
        assert "EIN" in types
        assert "SSN" in types

    def test_pii_clean_document_has_empty_redactions(self):
        from auditai_data_normalization.pii import scrub
        result = scrub("The audit was conducted in accordance with GAAS.")
        assert result.total_replacements() == 0

    def test_end_to_end_docx_no_pii_in_cleaned_text(self, tmp_path):
        """Full pipeline: DOCX with PII → normalized → no PII in output."""
        from auditai_data_normalization.normalize import normalize_document

        docx_path = tmp_path / "pii_test.docx"
        _make_simple_docx(docx_path, [
            "Organization: Harshwal Foundation LLC",
            "EIN: 94-9999999",
            "Completed by: MS1",
            "Statement of Financial Position Date: 06/30/2025",
        ])

        rec = normalize_document(str(docx_path), run_parallel=False)

        assert rec.pii_scrubbed is True
        # EIN must not appear in any text field
        assert "94-9999999" not in rec.cleaned_text
        for section in rec.sections:
            assert "94-9999999" not in section.content


# ---------------------------------------------------------------------------
# Additional: field_aliases resolution
# ---------------------------------------------------------------------------

class TestFieldAliases:
    """Spot-check that the alias file resolves real document labels."""

    @pytest.mark.parametrize("label,expected_canonical", [
        ("organization",                          "client_name"),
        ("statement of financial position date",  "fiscal_year_end"),
        ("completed by",                          "preparer_id"),
        ("engagement partner",                    "engagement_partner"),
        ("ep:",                                   "engagement_partner"),
        ("eng. partner",                          "engagement_partner"),
        ("single audit",                          "includes_single_audit"),
        ("total assets",                          "total_assets"),
        ("net assets without donor restrictions", "net_assets"),
        ("salaries and wages",                    "salaries_wages"),
        ("management and general",                "management_general_expenses"),
        ("opening_debit",                         "debit"),
        ("account_code",                          "account_code"),
    ])
    def test_alias_lookup(self, label: str, expected_canonical: str):
        from auditai_data_normalization.normalize import load_aliases
        aliases = load_aliases()
        result = aliases.get(label.lower().strip())
        assert result == expected_canonical, (
            f"'{label}' → '{result}', expected '{expected_canonical}'"
        )


# ---------------------------------------------------------------------------
# Additional: review_queue
# ---------------------------------------------------------------------------

class TestReviewQueue:
    """Spot-check review_queue enqueue / approve flow."""

    def test_enqueue_and_approve(self, tmp_path):
        from auditai_data_normalization.schema import DocumentRecord
        from auditai_data_normalization.review_queue import (
            enqueue, approve_record, pending, queue_stats,
        )

        queue_file = tmp_path / "queue.csv"

        rec = DocumentRecord(
            file_name="test.docx",
            file_type="docx",
            file_hash="a" * 64,
            extraction_confidence=0.6,
            extraction_status="success",
            needs_review=True,
            pii_scrubbed=True,
        )

        enqueue(rec, queue_file)
        assert len(pending(queue_file)) == 1

        rec = approve_record(rec, reviewer_id="SH", queue_path=queue_file)
        assert rec.auditor_approved is True
        assert rec.reviewer_id == "SH"
        assert rec.needs_review is False
        assert len(pending(queue_file)) == 0

        stats = queue_stats(queue_file)
        assert stats["total"] == 1
        assert stats["corrected"] == 1
        assert stats["pending"] == 0